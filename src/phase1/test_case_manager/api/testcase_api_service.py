#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Test Case API Service Module for the Watsonx IPG Testing platform.

This module provides RESTful API endpoints for managing test cases, including
creation, retrieval, updating, and deletion operations. It serves as the interface
between the Test Case Manager Module and external systems or UIs.

Updated to use PostgreSQL database storage instead of filesystem storage.
"""

import os
import logging
import json
from typing import Dict, List, Any, Tuple, Optional, Union
from datetime import datetime
import traceback
import tempfile
from io import BytesIO
import uuid

# Import Flask framework
from flask import Flask, request, jsonify, Response, send_file
from flask_cors import CORS
import pandas as pd
import psycopg2
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


# Import from src.common
from src.common.logging.log_utils import setup_logger
from src.common.exceptions.custom_exceptions import (
    TestCaseNotFoundError,
    MetadataError,
    VersionControlError,
    InvalidVersionError,
    DatabaseError
)

# Import from other modules
from src.phase1.test_case_manager.testcase_generator import TestCaseGenerator
from src.phase1.test_case_manager.testcase_refiner import LLMHelper
from src.phase1.test_case_manager.version_controller import VersionController
from src.phase1.test_case_manager.metadata_manager import MetadataManager


# Setup logger
logger = logging.getLogger(__name__)

# Create Flask app
app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "*"}})  # Enable CORS for all routes

# Initialize components
test_case_generator = None
test_case_refiner = None
version_controller = None
metadata_manager = None

def initialize_components():
    """Initialize all the components needed for the API service."""
    global test_case_generator, test_case_refiner, version_controller, metadata_manager
    
    try:
        # Initialize components with default configurations
        test_case_generator = TestCaseGenerator()
        test_case_refiner = LLMHelper()
        version_controller = VersionController()
        metadata_manager = MetadataManager()  # This now connects to PostgreSQL
        
        logger.info("All components initialized successfully")
        return True
    except Exception as e:
        logger.error(f"Failed to initialize components: {str(e)}")
        return False

# Error handler for API endpoints
def handle_error(error: Exception, status_code: int = 500) -> Tuple[Dict[str, Any], int]:
    """
    Create a standardized error response.
    
    Args:
        error (Exception): The exception that occurred.
        status_code (int, optional): HTTP status code. Defaults to 500.
        
    Returns:
        Tuple[Dict[str, Any], int]: Error response and status code.
    """
    error_type = error.__class__.__name__
    
    # Map certain exceptions to appropriate status codes
    if isinstance(error, TestCaseNotFoundError):
        status_code = 404  # Not Found
    elif isinstance(error, InvalidVersionError):
        status_code = 404  # Not Found
    elif isinstance(error, (MetadataError, VersionControlError, DatabaseError)):
        status_code = 400  # Bad Request
    
    # Create error response
    response = {
        "status": "error",
        "error": {
            "type": error_type,
            "message": str(error),
            "timestamp": datetime.now().isoformat()
        }
    }
    
    # Add stack trace in debug mode
    if app.debug:
        response["error"]["stack_trace"] = traceback.format_exc()
    
    return response, status_code

# Helper function to convert DataFrame to in-memory file
def df_to_excel_bytes(df: pd.DataFrame) -> bytes:
    """
    Convert a DataFrame to Excel file bytes.
    
    Args:
        df (pd.DataFrame): The DataFrame to convert.
        
    Returns:
        bytes: The Excel file as bytes.
    """
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)
    return output.getvalue()

# Helper function to create a temporary file
def create_temp_file(content: bytes, suffix: str = '.xlsx') -> str:
    """
    Create a temporary file with the given content.
    
    Args:
        content (bytes): The file content.
        suffix (str, optional): The file suffix. Defaults to '.xlsx'.
        
    Returns:
        str: The path to the temporary file.
    """
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
        temp_file.write(content)
        temp_path = temp_file.name
    return temp_path

# Helper function to convert test case DataFrame to a database-ready format
def prepare_test_case_for_db(test_case_df: pd.DataFrame, test_case_id: str = None) -> Dict[str, Any]:
    """
    Convert a test case DataFrame to a format suitable for database storage.
    
    Args:
        test_case_df (pd.DataFrame): The test case DataFrame.
        test_case_id (str, optional): The test case ID. If None, extracted from DataFrame.
        
    Returns:
        Dict[str, Any]: The test case data ready for database storage.
    """
    if test_case_df.empty:
        raise ValueError("Empty test case DataFrame cannot be prepared for database storage")
    
    # Extract basic test case info from first row
    first_row = test_case_df.iloc[0]
    
    # Get or generate test case ID
    if not test_case_id:
        test_case_id = first_row.get("TEST CASE NUMBER")
        if not test_case_id:
            test_case_id = f"TC-{uuid.uuid4().hex[:8].upper()}"
    
    # Extract other metadata
    metadata = {
        "TEST_CASE_ID": test_case_id,
        "TEST_CASE": first_row.get("TEST CASE", ""),
        "MODULE": first_row.get("SUBJECT", "Unknown"),
        "OWNER": first_row.get("TEST USER ID/ROLE", "Unassigned"),
        "TEST_TYPE": first_row.get("TYPE", "Functional"),
        "STATUS": "Draft",
        "PRIORITY": "Medium",
        "AUTOMATION_STATUS": "Manual",
        "CREATED_DATE": datetime.now(),
        "MODIFIED_DATE": datetime.now(),
    }
    
    # Convert DataFrame to Excel for storage
    file_content = df_to_excel_bytes(test_case_df)
    
    return {
        "metadata": metadata,
        "file_content": file_content,
        "file_name": f"{test_case_id}.xlsx",
        "file_type": "xlsx"
    }


# Health check endpoint
@app.route('/api/health', methods=['GET'])
def health_check():
    """
    Check the health of the API service.
    
    Returns:
        Tuple[Dict[str, Any], int]: Health status and HTTP status code.
    """
    # Check if all components are initialized
    is_healthy = (test_case_generator is not None and
                  test_case_refiner is not None and
                  version_controller is not None and
                  metadata_manager is not None)
    
    # Check database connection if metadata manager is initialized
    db_status = "unknown"
    if metadata_manager is not None:
        try:
            db_status = "connected" if metadata_manager.check_database_connection() else "disconnected"
        except Exception as e:
            logger.error(f"Database health check failed: {str(e)}")
            db_status = "error"
    
    response = {
        "status": "healthy" if (is_healthy and db_status == "connected") else "unhealthy",
        "timestamp": datetime.now().isoformat(),
        "components": {
            "test_case_generator": test_case_generator is not None,
            "test_case_refiner": test_case_refiner is not None,
            "version_controller": version_controller is not None,
            "metadata_manager": metadata_manager is not None,
            "database": db_status
        }
    }
    
    return response, 200 if (is_healthy and db_status == "connected") else 503


@app.route('/api/test-cases/generate', methods=['POST'])
def generate_test_case() -> Tuple[Dict[str, Any], int]:
    """
    Generate a test case from user input and store it in the database.
    
    Request body:
        - prompt (str): The user input for test case generation
        - output_path (str, optional): Legacy parameter, no longer used
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        # Extract the user input, checking both 'prompt' and 'scenario' fields for compatibility
        if 'prompt' in request_data:
            user_input = request_data['prompt']
        elif 'scenario' in request_data and isinstance(request_data['scenario'], dict):
            # If scenario is provided as a dict, try to extract description or name
            if 'description' in request_data['scenario']:
                user_input = request_data['scenario']['description']
            elif 'name' in request_data['scenario']:
                user_input = request_data['scenario']['name']
            else:
                # If no description or name, convert the entire scenario to a string
                user_input = str(request_data['scenario'])
        else:
            return {"status": "error", "message": "Missing prompt or scenario data"}, 400
        
        # Legacy output_path parameter, no longer used but kept for backward compatibility
        _ = request_data.get('output_path')
        
        # Generate test case directly from the prompt
        test_case_df = test_case_generator.generate_test_case_from_prompt(user_input)
        
        # Prepare test case for database storage
        test_case_data = prepare_test_case_for_db(test_case_df)
        
        # Store in database using metadata manager
        test_case_id = metadata_manager.create_test_case_metadata(
            test_case_data["metadata"], 
            created_by=request_data.get('created_by', 'API User')
        )
        
        # Store file content
        metadata_manager.store_test_case_file_content(
            test_case_id,
            test_case_data["file_name"],
            test_case_data["file_content"],
            test_case_data["file_type"],
            uploaded_by=request_data.get('created_by', 'API User')
        )
        
        # Return the test case data
        return {
            "status": "success",
            "message": "Test case generated successfully",
            "data": {
                "test_case_id": test_case_id,
                "test_case": test_case_df.to_dict('records'),
                "file_name": test_case_data["file_name"]
            }
        }, 200
        
    except Exception as e:
        return handle_error(e)


@app.route('/api/test-cases/generate-from-prompt', methods=['POST'])
def generate_test_case_from_prompt() -> Tuple[Dict[str, Any], int]:
    """
    Generate a test case from a simple text prompt and store it in the database.
    
    Request body:
        - prompt (str): Simple text prompt like "Generate login test cases for Admin user"
        - output_path (str, optional): Legacy parameter, no longer used
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'prompt' not in request_data:
            return {"status": "error", "message": "Missing prompt"}, 400
        
        prompt = request_data['prompt']
        
        # Legacy output_path parameter, no longer used but kept for backward compatibility
        _ = request_data.get('output_path')
        
        # Generate test case from prompt
        test_case_df = test_case_generator.generate_test_case_from_prompt(prompt)
        
        # Prepare test case for database storage
        test_case_data = prepare_test_case_for_db(test_case_df)
        
        # Store in database using metadata manager
        test_case_id = metadata_manager.create_test_case_metadata(
            test_case_data["metadata"], 
            created_by=request_data.get('created_by', 'API User')
        )
        
        # Store file content
        metadata_manager.store_test_case_file_content(
            test_case_id,
            test_case_data["file_name"],
            test_case_data["file_content"],
            test_case_data["file_type"],
            uploaded_by=request_data.get('created_by', 'API User')
        )
        
        # Return the test case data
        return {
            "status": "success",
            "message": "Test case generated successfully from prompt",
            "data": {
                "test_case_id": test_case_id,
                "test_case": test_case_df.to_dict('records'),
                "file_name": test_case_data["file_name"]
            }
        }, 200
        
    except Exception as e:
        # Check if it's our custom error type
        error_type = type(e).__name__
        error_message = str(e)
        logger.error(f"Test case generation failed: {error_message}")
        
        # Return a specific error response with the failure reason
        return {
            "status": "error",
            "message": "Failed to generate test case",
            "error_type": error_type,
            "details": error_message,
            "data": {
                "test_case": []  # Empty test case data
            }
        }, 400


@app.route('/api/test-cases/generate-batch', methods=['POST'])
def generate_test_cases_batch() -> Tuple[Dict[str, Any], int]:
    """
    Generate multiple test cases from scenarios and store them in the database.
    
    Request body:
        - scenarios (List[Dict]): List of test scenarios
        - output_dir (str, optional): Legacy parameter, no longer used
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'scenarios' not in request_data:
            return {"status": "error", "message": "Missing scenarios data"}, 400
        
        scenarios = request_data['scenarios']
        
        # Legacy output_dir parameter, no longer used but kept for backward compatibility
        _ = request_data.get('output_dir')
        
        # Generate test cases
        test_cases = test_case_generator.generate_test_cases_batch(scenarios)
        
        # Store each test case in the database
        test_case_ids = []
        for scenario_id, test_case_df in test_cases.items():
            # Prepare test case for database storage
            test_case_data = prepare_test_case_for_db(test_case_df)
            
            # Store in database using metadata manager
            test_case_id = metadata_manager.create_test_case_metadata(
                test_case_data["metadata"], 
                created_by=request_data.get('created_by', 'API User')
            )
            
            # Store file content
            metadata_manager.store_test_case_file_content(
                test_case_id,
                test_case_data["file_name"],
                test_case_data["file_content"],
                test_case_data["file_type"],
                uploaded_by=request_data.get('created_by', 'API User')
            )
            
            test_case_ids.append(test_case_id)
        
        # Return the result
        return {
            "status": "success",
            "message": f"Generated {len(test_cases)} test cases",
            "data": {
                "scenario_ids": list(test_cases.keys()),
                "test_case_ids": test_case_ids
            }
        }, 200
        
    except Exception as e:
        return handle_error(e)


@app.route('/api/test-cases/process-scenario-file', methods=['POST'])
def process_scenario_file() -> Tuple[Dict[str, Any], int]:
    """
    Process a scenario file to generate and store test cases in the database.
    
    Request body:
        - scenario_file_path (str): Path to the scenario file or file content
        - output_dir (str, optional): Legacy parameter, no longer used
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json or {}
        
        if 'scenario_file_path' not in request_data and 'scenario_file_content' not in request_data:
            return {"status": "error", "message": "Missing scenario_file_path or scenario_file_content"}, 400
        
        # Legacy output_dir parameter, no longer used but kept for backward compatibility
        _ = request_data.get('output_dir')
        
        # Process the scenario file
        if 'scenario_file_path' in request_data:
            # Legacy file path approach
            scenario_file_path = request_data['scenario_file_path']
            test_case_data_list = test_case_generator.process_scenario_file(scenario_file_path)
        else:
            # New approach with file content
            scenario_file_content = request_data['scenario_file_content']
            
            # Create a temporary file to process the content
            with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as temp_file:
                temp_file.write(scenario_file_content)
                temp_path = temp_file.name
            
            try:
                test_case_data_list = test_case_generator.process_scenario_file(temp_path)
            finally:
                # Clean up temporary file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
        
        # Store each test case in the database
        test_case_ids = []
        for test_case_data in test_case_data_list:
            # Extract test case DataFrame
            test_case_df = test_case_data.get('test_case_df')
            if test_case_df is None or test_case_df.empty:
                continue
            
            # Prepare test case for database storage
            db_test_case_data = prepare_test_case_for_db(test_case_df)
            
            # Store in database using metadata manager
            test_case_id = metadata_manager.create_test_case_metadata(
                db_test_case_data["metadata"], 
                created_by=request_data.get('created_by', 'API User')
            )
            
            # Store file content
            metadata_manager.store_test_case_file_content(
                test_case_id,
                db_test_case_data["file_name"],
                db_test_case_data["file_content"],
                db_test_case_data["file_type"],
                uploaded_by=request_data.get('created_by', 'API User')
            )
            
            test_case_ids.append(test_case_id)
        
        # Return the result
        return {
            "status": "success",
            "message": f"Generated {len(test_case_ids)} test case files",
            "data": {
                "test_case_ids": test_case_ids
            }
        }, 200
        
    except Exception as e:
        return handle_error(e)
    


# Version Controller Endpoints
@app.route('/api/test-cases/versions/check-in', methods=['POST'])
def check_in_new_version() -> Tuple[Dict[str, Any], int]:
    """
    Check in a new version of a test case.
    
    Request body:
        - test_case_id (str): The test case ID
        - test_case_content (bytes, optional): The new test case content
        - change_comment (str, optional): Comment describing the changes
        - changed_by (str, optional): Name/ID of the person who made the changes
        - notify_owner (bool, optional): Whether to notify the test case owner
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'test_case_id' not in request_data:
            return {"status": "error", "message": "Missing test_case_id"}, 400
        
        test_case_id = request_data['test_case_id']
        change_comment = request_data.get('change_comment')
        changed_by = request_data.get('changed_by')
        notify_owner = request_data.get('notify_owner', True)
        
        # Check if test case exists in database
        metadata = metadata_manager.get_test_case_metadata(test_case_id)
        if not metadata:
            return {"status": "error", "message": f"Test case {test_case_id} not found"}, 404
        
        # Get current version number
        current_version = metadata.get("VERSION", "1.0")
        
        # If test_case_content is provided, update the file content
        if 'test_case_content' in request_data:
            test_case_content = request_data['test_case_content']
            
            # Check if content is different from current version
            _, _, current_content = metadata_manager.retrieve_test_case_file_content(test_case_id)
            
            # Compare content (simple binary comparison)
            if current_content == test_case_content:
                # No changes, return early
                return {
                    "status": "success",
                    "message": "No changes detected",
                    "data": {
                        "test_case_id": test_case_id,
                        "version": current_version,
                        "is_new_version": False
                    }
                }, 200
            
            # Calculate new version number
            new_version = _increment_version(current_version)
            
            # Update metadata with new version
            metadata_manager.update_test_case_metadata(
                test_case_id,
                {
                    "VERSION": new_version,
                    "MODIFIED_DATE": datetime.now()
                },
                modified_by=changed_by
            )
            
            # Store the new version with version tag
            metadata_manager.store_test_case_file_content(
                test_case_id,
                f"{test_case_id}_v{new_version}.xlsx",
                test_case_content,
                "xlsx",
                uploaded_by=changed_by
            )
            
            # Store version history
            _store_version_history(
                test_case_id,
                current_version,
                new_version,
                change_comment,
                changed_by
            )
            
            # Notify owner if requested
            if notify_owner and metadata.get("OWNER"):
                # This would be implemented with a notification service
                # For now, just log the notification
                logger.info(f"Would notify owner {metadata.get('OWNER')} about new version {new_version} of test case {test_case_id}")
        
            # Return the result
            return {
                "status": "success",
                "message": "New version checked in successfully",
                "data": {
                    "test_case_id": test_case_id,
                    "version": new_version,
                    "previous_version": current_version,
                    "is_new_version": True,
                    "change_comment": change_comment
                }
            }, 200
        else:
            # No content provided, just return current version
            return {
                "status": "success",
                "message": "No new content provided",
                "data": {
                    "test_case_id": test_case_id,
                    "version": current_version,
                    "is_new_version": False
                }
            }, 200
        
    except Exception as e:
        return handle_error(e)

def _increment_version(version_str: str) -> str:
    """
    Increment a version string.
    
    Args:
        version_str (str): Current version string (e.g., "1.0").
        
    Returns:
        str: Incremented version string.
    """
    try:
        # Parse version string
        parts = version_str.split('.')
        if len(parts) == 1:
            # Single number, just increment
            return str(int(parts[0]) + 1)
        elif len(parts) == 2:
            # Major.minor format, increment minor
            major = int(parts[0])
            minor = int(parts[1]) + 1
            return f"{major}.{minor}"
        else:
            # More complex format, increment last part
            last_part = int(parts[-1]) + 1
            parts[-1] = str(last_part)
            return '.'.join(parts)
    except ValueError:
        # If parsing fails, just append ".1"
        return f"{version_str}.1"

def _store_version_history(test_case_id: str, old_version: str, new_version: str,
                          change_comment: str, changed_by: str) -> None:
    """
    Store version history in the database.
    
    Args:
        test_case_id (str): The test case ID.
        old_version (str): Previous version.
        new_version (str): New version.
        change_comment (str): Comment describing changes.
        changed_by (str): Person who made the changes.
    """
    # In a real implementation, this would store in a version history table
    # For now, we'll just use the metadata history which is already stored
    # when update_test_case_metadata is called
    pass

@app.route('/api/test-cases/<test_case_id>/versions', methods=['GET'])
def get_version_history(test_case_id: str) -> Tuple[Dict[str, Any], int]:
    """
    Get the version history for a test case.
    
    Args:
        test_case_id (str): The test case ID.
        
    Returns:
        Tuple[Dict[str, Any], int]: Version history and HTTP status code.
    """
    try:
        # Check if test case exists in database
        metadata = metadata_manager.get_test_case_metadata(test_case_id)
        if not metadata:
            return {"status": "error", "message": f"Test case {test_case_id} not found"}, 404
        
        # Get metadata history
        history = metadata_manager.get_metadata_history(test_case_id)
        
        # Filter for version changes
        version_history = []
        for entry in history:
            if entry["field_name"] == "VERSION":
                version_history.append({
                    "version": entry["new_value"],
                    "previous_version": entry["old_value"],
                    "changed_at": entry["changed_at"],
                    "changed_by": entry["changed_by"]
                })
        
        # Add current version if empty
        if not version_history:
            version_history.append({
                "version": metadata.get("VERSION", "1.0"),
                "previous_version": None,
                "changed_at": metadata.get("CREATED_DATE") or datetime.now().isoformat(),
                "changed_by": metadata.get("CREATED_BY", "System")
            })
        
        # Return the history
        return {
            "status": "success",
            "data": {
                "test_case_id": test_case_id,
                "current_version": metadata.get("VERSION", "1.0"),
                "version_history": version_history
            }
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/<test_case_id>/versions/<version>', methods=['GET'])
def get_test_case_version(test_case_id: str, version: str) -> Response:
    """
    Get a specific version of a test case.
    
    Args:
        test_case_id (str): The test case ID.
        version (str): The version to retrieve.
        
    Returns:
        Response: Excel file response or error.
    """
    try:
        # Check if test case exists in database
        metadata = metadata_manager.get_test_case_metadata(test_case_id)
        if not metadata:
            return jsonify({"status": "error", "message": f"Test case {test_case_id} not found"}), 404
        
        # Get the test case file content from the database
        # In a real implementation, this would retrieve a specific version from a version table
        # For now, we can only get the current version
        file_name, file_type, file_content = metadata_manager.retrieve_test_case_file_content(test_case_id)
        
        # Check if this is the current version
        current_version = metadata.get("VERSION", "1.0")
        if version != current_version:
            # For historical versions, we would need to retrieve from a version history table
            # For now, just return an error
            return jsonify({"status": "error", "message": f"Version {version} not available"}), 404
        
        # Generate appropriate file name
        download_name = f"{test_case_id}_v{version}.xlsx"
        
        # Return the file
        return send_file(
            BytesIO(file_content),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=download_name
        )
        
    except Exception as e:
        error_response, status_code = handle_error(e)
        return jsonify(error_response), status_code

@app.route('/api/test-cases/<test_case_id>/versions/compare', methods=['GET'])
def compare_versions(test_case_id: str) -> Tuple[Dict[str, Any], int]:
    """
    Compare two versions of a test case.
    
    Args:
        test_case_id (str): The test case ID.
        
    Query parameters:
        - version1 (str): First version to compare.
        - version2 (str): Second version to compare.
        
    Returns:
        Tuple[Dict[str, Any], int]: Comparison results and HTTP status code.
    """
    try:
        # Get query parameters
        version1 = request.args.get('version1')
        version2 = request.args.get('version2')
        
        if not version1 or not version2:
            return {"status": "error", "message": "Missing version parameters"}, 400
        
        # Check if test case exists in database
        metadata = metadata_manager.get_test_case_metadata(test_case_id)
        if not metadata:
            return {"status": "error", "message": f"Test case {test_case_id} not found"}, 404
        
        # In a real implementation, this would compare two versions from a version history table
        # For now, just return a placeholder result
        comparison = {
            "test_case_id": test_case_id,
            "version1": version1,
            "version2": version2,
            "changes": [
                {
                    "type": "info",
                    "message": "Version comparison not implemented yet. Only the current version is available."
                }
            ]
        }
        
        # Return the comparison results
        return {
            "status": "success",
            "data": comparison
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/<test_case_id>/status', methods=['PUT'])
def update_test_case_status(test_case_id: str) -> Tuple[Dict[str, Any], int]:
    """
    Update the status of a test case.
    
    Args:
        test_case_id (str): The test case ID.
        
    Request body:
        - status (str): The new status ("Under Maintenance", "Active", "Obsolete", etc.).
        - changed_by (str, optional): Person making the update.
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'status' not in request_data:
            return {"status": "error", "message": "Missing status"}, 400
        
        new_status = request_data['status']
        changed_by = request_data.get('changed_by')
        
        # Check if test case exists in database
        metadata = metadata_manager.get_test_case_metadata(test_case_id)
        if not metadata:
            return {"status": "error", "message": f"Test case {test_case_id} not found"}, 404
        
        # Update status
        updated_metadata = metadata_manager.update_test_case_status(
            test_case_id,
            new_status,
            changed_by
        )
        
        # Return the result
        return {
            "status": "success",
            "message": f"Test case status updated to {new_status}",
            "data": {
                "test_case_id": test_case_id,
                "previous_status": metadata.get("STATUS"),
                "new_status": new_status,
                "timestamp": datetime.now().isoformat()
            }
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/<test_case_id>/sharepoint', methods=['POST'])
def upload_to_sharepoint(test_case_id: str) -> Tuple[Dict[str, Any], int]:
    """
    Upload a test case version to SharePoint.
    
    Args:
        test_case_id (str): The test case ID.
        
    Request body:
        - version (str, optional): The version to upload.
        - sharepoint_folder (str, optional): The SharePoint folder path.
        
    Returns:
        Tuple[Dict[str, Any], int]: Upload result and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json or {}
        
        version = request_data.get('version')
        sharepoint_folder = request_data.get('sharepoint_folder')
        
        # Check if test case exists in database
        metadata = metadata_manager.get_test_case_metadata(test_case_id)
        if not metadata:
            return {"status": "error", "message": f"Test case {test_case_id} not found"}, 404
        
        # Get the file content from the database
        file_name, file_type, file_content = metadata_manager.retrieve_test_case_file_content(test_case_id)
        
        # In a real implementation, this would upload to SharePoint
        # For now, just log the action
        logger.info(f"Would upload test case {test_case_id} to SharePoint folder {sharepoint_folder}")
        
        # Return the result
        return {
            "status": "success",
            "message": f"Test case uploaded to SharePoint successfully",
            "data": {
                "test_case_id": test_case_id,
                "version": version or metadata.get("VERSION", "1.0"),
                "sharepoint_folder": sharepoint_folder,
                "file_name": file_name,
                "upload_date": datetime.now().isoformat()
            }
        }, 200
        
    except Exception as e:
        return handle_error(e)

# Metadata Manager Endpoints
@app.route('/api/test-cases/metadata/<test_case_id>', methods=['GET'])
def get_test_case_metadata(test_case_id: str) -> Tuple[Dict[str, Any], int]:
    """
    Get metadata for a test case.
    
    Args:
        test_case_id (str): The test case ID.
        
    Returns:
        Tuple[Dict[str, Any], int]: Metadata and HTTP status code.
    """
    try:
        # Get metadata
        metadata = metadata_manager.get_test_case_metadata(test_case_id)
        
        if not metadata:
            return {"status": "error", "message": f"Test case {test_case_id} not found"}, 404
        
        # Check if file exists
        file_exists = metadata_manager.file_exists_for_test_case(test_case_id)
        
        # Return the metadata
        return {
            "status": "success",
            "data": {
                "metadata": metadata,
                "file_exists": file_exists
            }
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/metadata', methods=['POST'])
def create_test_case_metadata() -> Tuple[Dict[str, Any], int]:
    """
    Create metadata for a new test case.
    
    Request body:
        - test_case_data (Dict): Basic test case data.
        - created_by (str, optional): Person creating the test case.
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'test_case_data' not in request_data:
            return {"status": "error", "message": "Missing test_case_data"}, 400
        
        test_case_data = request_data['test_case_data']
        created_by = request_data.get('created_by')
        
        # Create metadata
        test_case_id = metadata_manager.create_test_case_metadata(
            test_case_data,
            created_by=created_by
        )
        
        # Get the created metadata
        metadata = metadata_manager.get_test_case_metadata(test_case_id)
        
        # Return the result
        return {
            "status": "success",
            "message": f"Metadata created for test case {test_case_id}",
            "data": metadata
        }, 201
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/metadata/<test_case_id>', methods=['PUT'])
def update_test_case_metadata(test_case_id: str) -> Tuple[Dict[str, Any], int]:
    """
    Update metadata for an existing test case.
    
    Args:
        test_case_id (str): The test case ID.
        
    Request body:
        - updates (Dict): Metadata fields to update.
        - modified_by (str, optional): Person making the updates.
        
    Returns:
        Tuple[Dict[str, Any], int]: Updated metadata and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'updates' not in request_data:
            return {"status": "error", "message": "Missing updates"}, 400
        
        updates = request_data['updates']
        modified_by = request_data.get('modified_by')
        
        # Update metadata
        updated_metadata = metadata_manager.update_test_case_metadata(
            test_case_id,
            updates,
            modified_by=modified_by
        )
        
        # Return the updated metadata
        return {
            "status": "success",
            "message": f"Metadata updated for test case {test_case_id}",
            "data": updated_metadata
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/metadata/<test_case_id>', methods=['DELETE'])
def delete_test_case_metadata(test_case_id: str) -> Tuple[Dict[str, Any], int]:
    """
    Delete metadata for a test case.
    
    Args:
        test_case_id (str): The test case ID.
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code.
    """
    try:
        # Delete metadata
        success = metadata_manager.delete_test_case_metadata(test_case_id)
        
        if not success:
            return {"status": "error", "message": f"Failed to delete metadata for test case {test_case_id}"}, 500
        
        # Return the result
        return {
            "status": "success",
            "message": f"Metadata deleted for test case {test_case_id}"
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/search', methods=['POST'])
def search_test_cases() -> Tuple[Dict[str, Any], int]:
    """
    Search for test cases based on metadata criteria.
    
    Request body:
        - criteria (Dict): Search criteria.
        
    Returns:
        Tuple[Dict[str, Any], int]: Search results and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'criteria' not in request_data:
            return {"status": "error", "message": "Missing search criteria"}, 400
        
        criteria = request_data['criteria']
        
        # Search test cases
        results = metadata_manager.search_test_cases(criteria)
        
        # Return the results
        return {
            "status": "success",
            "message": f"Found {len(results)} matching test cases",
            "data": results
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/<test_case_id>/ownership', methods=['PUT'])
def update_test_case_owner(test_case_id: str) -> Tuple[Dict[str, Any], int]:
    """
    Update the owner of a test case.
    
    Args:
        test_case_id (str): The test case ID.
        
    Request body:
        - owner (str): The new owner.
        - modified_by (str, optional): Person making the update.
        
    Returns:
        Tuple[Dict[str, Any], int]: Updated metadata and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'owner' not in request_data:
            return {"status": "error", "message": "Missing owner"}, 400
        
        new_owner = request_data['owner']
        modified_by = request_data.get('modified_by')
        
        # Update owner
        updated_metadata = metadata_manager.update_test_case_owner(
            test_case_id,
            new_owner,
            modified_by=modified_by
        )
        
        # Return the updated metadata
        return {
            "status": "success",
            "message": f"Owner updated for test case {test_case_id}",
            "data": updated_metadata
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/<test_case_id>/automation-status', methods=['PUT'])
def update_automation_status(test_case_id: str) -> Tuple[Dict[str, Any], int]:
    """
    Update the automation status of a test case.
    
    Args:
        test_case_id (str): The test case ID.
        
    Request body:
        - status (str): The new automation status.
        - modified_by (str, optional): Person making the update.
        
    Returns:
        Tuple[Dict[str, Any], int]: Updated metadata and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'status' not in request_data:
            return {"status": "error", "message": "Missing status"}, 400
        
        new_status = request_data['status']
        modified_by = request_data.get('modified_by')
        
        # Update automation status
        updated_metadata = metadata_manager.update_automation_status(
            test_case_id,
            new_status,
            modified_by=modified_by
        )
        
        # Return the updated metadata
        return {
            "status": "success",
            "message": f"Automation status updated for test case {test_case_id}",
            "data": updated_metadata
        }, 200
        
    except Exception as e:
        return handle_error(e)


@app.route('/api/test-cases/<test_case_id>/execution-result', methods=['PUT'])
def update_execution_result(test_case_id: str) -> Tuple[Dict[str, Any], int]:
    """
    Update the execution result for a test case.
    
    Args:
        test_case_id (str): The test case ID.
        
    Request body:
        - result (str): The execution result.
        - executed_by (str, optional): Person who executed the test.
        
    Returns:
        Tuple[Dict[str, Any], int]: Updated metadata and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'result' not in request_data:
            return {"status": "error", "message": "Missing result"}, 400
        
        result = request_data['result']
        executed_by = request_data.get('executed_by')
        
        # Update execution result
        updated_metadata = metadata_manager.update_test_execution_result(
            test_case_id,
            result,
            executed_by=executed_by
        )
        
        # Return the updated metadata
        return {
            "status": "success",
            "message": f"Execution result updated for test case {test_case_id}",
            "data": updated_metadata
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/<test_case_id>/tags', methods=['POST'])
def add_tags_to_test_case(test_case_id: str) -> Tuple[Dict[str, Any], int]:
    """
    Add tags to a test case.
    
    Args:
        test_case_id (str): The test case ID.
        
    Request body:
        - tags (List[str]): Tags to add.
        - modified_by (str, optional): Person making the update.
        
    Returns:
        Tuple[Dict[str, Any], int]: Updated metadata and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'tags' not in request_data:
            return {"status": "error", "message": "Missing tags"}, 400
        
        tags = request_data['tags']
        modified_by = request_data.get('modified_by')
        
        # Add tags
        updated_metadata = metadata_manager.add_tags_to_test_case(
            test_case_id,
            tags,
            modified_by=modified_by
        )
        
        # Return the updated metadata
        return {
            "status": "success",
            "message": f"Tags added to test case {test_case_id}",
            "data": updated_metadata
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/<test_case_id>/tags', methods=['DELETE'])
def remove_tags_from_test_case(test_case_id: str) -> Tuple[Dict[str, Any], int]:
    """
    Remove tags from a test case.
    
    Args:
        test_case_id (str): The test case ID.
        
    Request body:
        - tags (List[str]): Tags to remove.
        - modified_by (str, optional): Person making the update.
        
    Returns:
        Tuple[Dict[str, Any], int]: Updated metadata and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'tags' not in request_data:
            return {"status": "error", "message": "Missing tags"}, 400
        
        tags = request_data['tags']
        modified_by = request_data.get('modified_by')
        
        # Remove tags
        updated_metadata = metadata_manager.remove_tags_from_test_case(
            test_case_id,
            tags,
            modified_by=modified_by
        )
        
        # Return the updated metadata
        return {
            "status": "success",
            "message": f"Tags removed from test case {test_case_id}",
            "data": updated_metadata
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/tags', methods=['GET'])
def get_all_tags() -> Tuple[Dict[str, Any], int]:
    """
    Get all tags used across test cases.
    
    Returns:
        Tuple[Dict[str, Any], int]: List of tags and HTTP status code.
    """
    try:
        # Get all tags
        tags = metadata_manager.get_all_tags()
        
        # Return the tags
        return {
            "status": "success",
            "data": tags
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/file-content/<test_case_id>', methods=['GET'])
def get_test_case_file_content(test_case_id: str) -> Response:
    """
    Get the file content for a test case.
    
    Args:
        test_case_id (str): The test case ID.
        
    Query parameters:
        - format (str, optional): Output format (excel, json, binary). Default: excel.
        
    Returns:
        Response: File content response.
    """
    try:
        output_format = request.args.get('format', 'excel').lower()
        
        # Check if test case exists in database
        metadata = metadata_manager.get_test_case_metadata(test_case_id)
        if not metadata:
            return jsonify({"status": "error", "message": f"Test case {test_case_id} not found"}), 404
        
        # Get the file content from the database
        file_name, file_type, file_content = metadata_manager.retrieve_test_case_file_content(test_case_id)
        
        if not file_content:
            return jsonify({"status": "error", "message": f"No file content found for test case {test_case_id}"}), 404
        
        if output_format == 'binary':
            # Return raw binary content
            return Response(
                file_content,
                mimetype='application/octet-stream',
                headers={'Content-Disposition': f'attachment; filename="{file_name}"'}
            )
        
        elif output_format == 'excel':
            # Return as Excel file
            return send_file(
                BytesIO(file_content),
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=file_name
            )
        
        elif output_format == 'json':
            # Try to parse Excel content and return as JSON
            try:
                df = pd.read_excel(BytesIO(file_content), engine='openpyxl')
                return jsonify({
                    "status": "success",
                    "data": df.to_dict('records')
                })
            except Exception as e:
                logger.error(f"Failed to convert Excel to JSON: {str(e)}")
                return jsonify({"status": "error", "message": f"Failed to convert to JSON: {str(e)}"}), 500
        
        else:
            return jsonify({"status": "error", "message": f"Unsupported format: {output_format}"}), 400
        
    except Exception as e:
        error_response, status_code = handle_error(e)
        return jsonify(error_response), status_code

@app.route('/api/test-cases/file-content/<test_case_id>', methods=['POST'])
def upload_test_case_file_content(test_case_id: str) -> Tuple[Dict[str, Any], int]:
    """
    Upload file content for a test case.
    
    Args:
        test_case_id (str): The test case ID.
        
    Request body:
        - file_content (bytes): The file content (for JSON requests).
        - file_name (str, optional): The file name.
        - file_type (str, optional): The file type.
        - uploaded_by (str, optional): Person uploading the file.
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code.
    """
    try:
        # Check if test case exists in database
        metadata = metadata_manager.get_test_case_metadata(test_case_id)
        if not metadata:
            return {"status": "error", "message": f"Test case {test_case_id} not found"}, 404
        
        file_content = None
        file_name = None
        file_type = None
        
        # Handle file upload or JSON content
        if request.files and 'file' in request.files:
            # Get file from request
            file = request.files['file']
            file_name = file.filename
            file_type = os.path.splitext(file_name)[1].lstrip('.')
            file_content = file.read()
        
        elif request.json and 'file_content' in request.json:
            # Get file content from JSON
            file_content = request.json['file_content']
            
            # Convert base64 if needed
            if isinstance(file_content, str):
                import base64
                try:
                    file_content = base64.b64decode(file_content)
                except Exception:
                    # Not base64, treat as raw bytes
                    file_content = file_content.encode('utf-8')
            
            file_name = request.json.get('file_name', f"{test_case_id}.xlsx")
            file_type = request.json.get('file_type', os.path.splitext(file_name)[1].lstrip('.') or 'xlsx')
        
        else:
            return {"status": "error", "message": "No file content provided"}, 400
        
        # Get uploader
        uploaded_by = request.form.get('uploaded_by') if request.form else request.json.get('uploaded_by') if request.json else None
        
        # Store in database
        metadata_manager.store_test_case_file_content(
            test_case_id,
            file_name,
            file_content,
            file_type,
            uploaded_by=uploaded_by
        )
        
        # Return success
        return {
            "status": "success",
            "message": f"File content uploaded for test case {test_case_id}",
            "data": {
                "test_case_id": test_case_id,
                "file_name": file_name,
                "file_type": file_type,
                "content_length": len(file_content),
                "upload_timestamp": datetime.now().isoformat()
            }
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/file-content/<test_case_id>', methods=['DELETE'])
def delete_test_case_file_content(test_case_id: str) -> Tuple[Dict[str, Any], int]:
    """
    Delete file content for a test case.
    
    Args:
        test_case_id (str): The test case ID.
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code.
    """
    try:
        # Check if test case exists in database
        metadata = metadata_manager.get_test_case_metadata(test_case_id)
        if not metadata:
            return {"status": "error", "message": f"Test case {test_case_id} not found"}, 404
        
        # Delete file content
        # In a real implementation, metadata_manager would have a method to delete file content
        # For now, simulate with a placeholder response
        logger.info(f"Would delete file content for test case {test_case_id}")
        
        # Update metadata to reflect deleted file
        metadata_manager.update_test_case_metadata(
            test_case_id,
            {
                "FILE_NAME": None,
                "FILE_TYPE": None,
                "MODIFIED_DATE": datetime.now()
            }
        )
        
        # Return success
        return {
            "status": "success",
            "message": f"File content deleted for test case {test_case_id}"
        }, 200
        
    except Exception as e:
        return handle_error(e)
    


@app.route('/api/test-cases/metadata/export', methods=['POST'])
def export_metadata() -> Tuple[Dict[str, Any], int]:
    """
    Export test case metadata to a JSON file stored in memory.
    
    Request body:
        - test_case_ids (List[str], optional): Specific test cases to export.
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code or file download.
    """
    try:
        # Get request data
        request_data = request.json or {}
        test_case_ids = request_data.get('test_case_ids')
        include_file_content = request_data.get('include_file_content', False)
        
        # Create a temporary file for JSON export
        with tempfile.NamedTemporaryFile(suffix='.json', delete=False) as temp_file:
            temp_path = temp_file.name
        
        try:
            # Export metadata to the temporary file
            count = metadata_manager.export_metadata_to_json(temp_path, test_case_ids)
            
            # Read the exported file
            with open(temp_path, 'r') as f:
                metadata_json = json.loads(f.read())
            
            # If requested, include file content
            if include_file_content and test_case_ids:
                for test_case_id in test_case_ids:
                    try:
                        _, file_type, file_content = metadata_manager.retrieve_test_case_file_content(test_case_id)
                        if file_content:
                            # Encode file content as base64
                            import base64
                            encoded_content = base64.b64encode(file_content).decode('utf-8')
                            
                            # Find the metadata entry for this test case
                            for entry in metadata_json:
                                if entry.get("TEST_CASE_ID") == test_case_id:
                                    # Add file content
                                    entry["FILE_CONTENT_BASE64"] = encoded_content
                                    break
                    except Exception as file_error:
                        logger.error(f"Failed to include file content for {test_case_id}: {str(file_error)}")
            
            # For API response
            download_format = request.args.get('download', 'false').lower()
            
            if download_format == 'true' or download_format == 'json':
                # Return as downloadable JSON file
                return send_file(
                    BytesIO(json.dumps(metadata_json, indent=2).encode('utf-8')),
                    mimetype='application/json',
                    as_attachment=True,
                    download_name=f"test_case_metadata_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                )
            else:
                # Return JSON in API response
                return {
                    "status": "success",
                    "message": f"Exported {count} test case metadata records",
                    "data": {
                        "count": count,
                        "metadata": metadata_json
                    }
                }, 200
                
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/metadata/import', methods=['POST'])
def import_metadata() -> Tuple[Dict[str, Any], int]:
    """
    Import test case metadata from a JSON file or content.
    
    Request body:
        - metadata_json (List[Dict], optional): Metadata as JSON.
        - overwrite (bool, optional): Whether to overwrite existing metadata.
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code.
    """
    try:
        metadata_json = None
        
        # Check if we have a file upload
        if request.files and 'metadata_file' in request.files:
            file = request.files['metadata_file']
            # Read file content
            file_content = file.read()
            metadata_json = json.loads(file_content)
        
        # Check if we have JSON data in request
        elif request.json and 'metadata_json' in request.json:
            metadata_json = request.json['metadata_json']
        
        else:
            return {"status": "error", "message": "Missing metadata_file or metadata_json"}, 400
        
        # Get overwrite flag
        overwrite = request.json.get('overwrite', False) if request.json else request.form.get('overwrite', 'false').lower() == 'true'
        include_file_content = request.json.get('include_file_content', True) if request.json else request.form.get('include_file_content', 'true').lower() == 'true'
        
        # Create a temporary file for the JSON data
        with tempfile.NamedTemporaryFile(suffix='.json', delete=False) as temp_file:
            temp_file.write(json.dumps(metadata_json).encode('utf-8'))
            temp_path = temp_file.name
        
        try:
            # Import metadata from the temporary file
            count = metadata_manager.import_metadata_from_json(temp_path, overwrite)
            
            # If file content is included, process it
            if include_file_content:
                file_count = 0
                for entry in metadata_json:
                    test_case_id = entry.get("TEST_CASE_ID")
                    if not test_case_id:
                        continue
                    
                    # Check if we have base64 encoded file content
                    if "FILE_CONTENT_BASE64" in entry:
                        try:
                            # Decode file content
                            import base64
                            file_content = base64.b64decode(entry["FILE_CONTENT_BASE64"])
                            
                            # Get file name and type
                            file_name = entry.get("FILE_NAME", f"{test_case_id}.xlsx")
                            file_type = entry.get("FILE_TYPE", os.path.splitext(file_name)[1].lstrip('.') or 'xlsx')
                            
                            # Store file content
                            metadata_manager.store_test_case_file_content(
                                test_case_id,
                                file_name,
                                file_content,
                                file_type,
                                uploaded_by="Import Process"
                            )
                            
                            file_count += 1
                        except Exception as file_error:
                            logger.error(f"Failed to import file content for {test_case_id}: {str(file_error)}")
            
            # Return the result
            return {
                "status": "success",
                "message": f"Imported {count} test case metadata records",
                "data": {
                    "metadata_count": count,
                    "file_count": file_count if include_file_content else 0
                }
            }, 200
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/bulk-export', methods=['POST'])
def bulk_export_test_cases() -> Tuple[Dict[str, Any], int]:
    """
    Export multiple test cases to a ZIP file.
    
    Request body:
        - test_case_ids (List[str]): List of test case IDs to export.
        - include_metadata (bool, optional): Whether to include metadata in export.
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code or file download.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'test_case_ids' not in request_data:
            return {"status": "error", "message": "Missing test_case_ids"}, 400
        
        test_case_ids = request_data['test_case_ids']
        include_metadata = request_data.get('include_metadata', True)
        
        # Create a temporary directory for files
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Create individual files for each test case
            exported_files = []
            
            for test_case_id in test_case_ids:
                try:
                    # Get metadata
                    metadata = metadata_manager.get_test_case_metadata(test_case_id)
                    if not metadata:
                        logger.warning(f"Test case not found: {test_case_id}")
                        continue
                    
                    # Get file content
                    file_name, file_type, file_content = metadata_manager.retrieve_test_case_file_content(test_case_id)
                    
                    if not file_content:
                        logger.warning(f"No file content found for test case: {test_case_id}")
                        continue
                    
                    # Create file path
                    file_path = os.path.join(temp_dir, file_name or f"{test_case_id}.xlsx")
                    
                    # Write to file
                    with open(file_path, 'wb') as f:
                        f.write(file_content)
                    
                    exported_files.append(file_path)
                    
                except Exception as case_error:
                    logger.error(f"Failed to export test case {test_case_id}: {str(case_error)}")
                    # Continue with next test case
            
            # Export metadata if requested
            if include_metadata:
                metadata_path = os.path.join(temp_dir, "metadata.json")
                metadata_manager.export_metadata_to_json(metadata_path, test_case_ids)
                exported_files.append(metadata_path)
            
            # Create ZIP file
            zip_file_path = os.path.join(temp_dir, f"test_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip")
            
            import zipfile
            with zipfile.ZipFile(zip_file_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_path in exported_files:
                    zipf.write(file_path, os.path.basename(file_path))
            
            # Check if we want a download or API response
            download_format = request.args.get('download', 'false').lower()
            
            if download_format == 'true' or download_format == 'zip':
                # Return ZIP file for download
                with open(zip_file_path, 'rb') as f:
                    zip_content = f.read()
                
                return send_file(
                    BytesIO(zip_content),
                    mimetype='application/zip',
                    as_attachment=True,
                    download_name=os.path.basename(zip_file_path)
                )
            else:
                # Return result as JSON
                return {
                    "status": "success",
                    "message": f"Exported {len(exported_files)} files for {len(test_case_ids)} test cases",
                    "data": {
                        "test_case_count": len(test_case_ids),
                        "file_count": len(exported_files),
                        "exported_test_cases": [tc_id for tc_id in test_case_ids if metadata_manager.get_test_case_metadata(tc_id)]
                    }
                }, 200
                
        finally:
            # Clean up temporary directory
            import shutil
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/test-cases/bulk-import', methods=['POST'])
def bulk_import_test_cases() -> Tuple[Dict[str, Any], int]:
    """
    Import multiple test cases from files or a ZIP archive.
    
    Request body for multipart/form-data:
        - files: List of files to import.
        - zip_file: ZIP archive containing files to import.
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code.
    """
    try:
        # Check if we have files
        files = []
        test_case_ids = []
        
        # Handle individual files
        if 'files[]' in request.files:
            file_list = request.files.getlist('files[]')
            
            for file in file_list:
                if file and file.filename:
                    # Create a temporary file
                    with tempfile.NamedTemporaryFile(suffix=os.path.splitext(file.filename)[1], delete=False) as temp_file:
                        file.save(temp_file.name)
                        files.append(temp_file.name)
        
        # Handle ZIP archive
        elif 'zip_file' in request.files:
            zip_file = request.files['zip_file']
            
            if zip_file and zip_file.filename:
                # Create a temporary file for the ZIP
                zip_temp = tempfile.NamedTemporaryFile(suffix='.zip', delete=False)
                zip_path = zip_temp.name
                zip_file.save(zip_path)
                
                # Create a temporary directory for extracted files
                temp_dir = tempfile.mkdtemp()
                
                try:
                    # Extract ZIP
                    import zipfile
                    with zipfile.ZipFile(zip_path, 'r') as zipf:
                        zipf.extractall(temp_dir)
                    
                    # Process each file
                    for root, _, file_names in os.walk(temp_dir):
                        for file_name in file_names:
                            file_path = os.path.join(root, file_name)
                            files.append(file_path)
                    
                finally:
                    # Clean up ZIP file
                    if os.path.exists(zip_path):
                        os.remove(zip_path)
        
        else:
            return {"status": "error", "message": "No files or ZIP archive provided"}, 400
        
        # Process each file
        try:
            # Get uploader from form or JSON
            uploaded_by = request.form.get('uploaded_by') if request.form else request.json.get('uploaded_by') if request.json else None
            
            for file_path in files:
                try:
                    # Import based on file extension
                    file_ext = os.path.splitext(file_path)[1].lower()
                    
                    if file_ext in ['.xlsx', '.xls']:
                        # Import Excel as test case
                        test_case_id = metadata_manager.import_test_case_from_excel(
                            file_path,
                            uploaded_by=uploaded_by
                        )
                        test_case_ids.append(test_case_id)
                    
                    elif file_ext == '.json':
                        # Check if it's metadata JSON
                        with open(file_path, 'r') as f:
                            try:
                                json_data = json.load(f)
                                
                                if isinstance(json_data, list) and all('TEST_CASE_ID' in item for item in json_data):
                                    # Import as metadata
                                    count = metadata_manager.import_metadata_from_json(file_path, overwrite=True)
                                    logger.info(f"Imported {count} test case records from {file_path}")
                            except json.JSONDecodeError:
                                logger.warning(f"Invalid JSON file: {file_path}")
                    
                    else:
                        logger.warning(f"Unsupported file type: {file_ext}")
                
                except Exception as file_error:
                    logger.error(f"Failed to import file {file_path}: {str(file_error)}")
                    # Continue with next file
        
        finally:
            # Clean up temporary files
            for file_path in files:
                if os.path.exists(file_path) and file_path.startswith(tempfile.gettempdir()):
                    os.remove(file_path)
            
            # Clean up temp directory if created
            if 'temp_dir' in locals() and os.path.exists(temp_dir):
                import shutil
                shutil.rmtree(temp_dir)
        
        # Return the result
        return {
            "status": "success",
            "message": f"Imported {len(test_case_ids)} test cases",
            "data": {
                "imported_test_case_ids": test_case_ids,
                "file_count": len(files)
            }
        }, 200
        
    except Exception as e:
        return handle_error(e)
    

@app.route('/api/stats', methods=['GET'])
def get_system_stats() -> Tuple[Dict[str, Any], int]:
    """
    Get statistics about test cases and database usage.
    
    Returns:
        Tuple[Dict[str, Any], int]: Statistics and HTTP status code.
    """
    try:
        # Get basic test case stats
        test_case_stats = metadata_manager.get_stats()
        
        # Get database stats if available
        try:
            db_stats = metadata_manager.get_database_stats()
        except Exception as db_error:
            logger.error(f"Failed to get database stats: {str(db_error)}")
            db_stats = {"error": str(db_error)}
        
        # Check database connection
        try:
            db_connection = metadata_manager.check_database_connection()
        except Exception as conn_error:
            logger.error(f"Failed to check database connection: {str(conn_error)}")
            db_connection = False
        
        # Return combined stats
        return {
            "status": "success",
            "data": {
                "test_cases": test_case_stats,
                "database": db_stats,
                "connection_status": "connected" if db_connection else "disconnected",
                "timestamp": datetime.now().isoformat()
            }
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/reports/test-cases', methods=['GET'])
def get_test_case_report() -> Tuple[Dict[str, Any], int]:
    """
    Get a report of test cases with various filtering options.
    
    Query parameters:
        - status (str, optional): Filter by status.
        - owner (str, optional): Filter by owner.
        - module (str, optional): Filter by module.
        - automation_status (str, optional): Filter by automation status.
        - created_after (str, optional): Filter by creation date.
        - created_before (str, optional): Filter by creation date.
        - modified_after (str, optional): Filter by modification date.
        - modified_before (str, optional): Filter by modification date.
        - limit (int, optional): Limit number of results.
        - offset (int, optional): Offset for pagination.
        - format (str, optional): Output format (json, excel, pdf).
        
    Returns:
        Tuple[Dict[str, Any], int]: Report data and HTTP status code.
    """
    try:
        # Collect filter criteria
        criteria = {}
        
        # Add filters from query parameters
        if 'status' in request.args:
            criteria["STATUS"] = request.args.get('status')
        
        if 'owner' in request.args:
            criteria["OWNER"] = request.args.get('owner')
        
        if 'module' in request.args:
            criteria["MODULE"] = request.args.get('module')
        
        if 'automation_status' in request.args:
            criteria["AUTOMATION_STATUS"] = request.args.get('automation_status')
        
        # Date-based filters need special handling
        if 'created_after' in request.args:
            criteria["CREATED_DATE"] = {"op": ">", "value": request.args.get('created_after')}
        
        if 'created_before' in request.args:
            if "CREATED_DATE" in criteria:
                # Update existing criteria to a between operation
                created_after = criteria["CREATED_DATE"]["value"]
                criteria["CREATED_DATE"] = {"op": "between", "value": [created_after, request.args.get('created_before')]}
            else:
                criteria["CREATED_DATE"] = {"op": "<", "value": request.args.get('created_before')}
        
        if 'modified_after' in request.args:
            criteria["MODIFIED_DATE"] = {"op": ">", "value": request.args.get('modified_after')}
        
        if 'modified_before' in request.args:
            if "MODIFIED_DATE" in criteria:
                # Update existing criteria to a between operation
                modified_after = criteria["MODIFIED_DATE"]["value"]
                criteria["MODIFIED_DATE"] = {"op": "between", "value": [modified_after, request.args.get('modified_before')]}
            else:
                criteria["MODIFIED_DATE"] = {"op": "<", "value": request.args.get('modified_before')}
        
        # Search test cases with criteria
        results = metadata_manager.search_test_cases(criteria)
        
        # Apply pagination if requested
        limit = request.args.get('limit', type=int)
        offset = request.args.get('offset', type=int, default=0)
        
        if limit is not None:
            results = results[offset:offset+limit]
        
        # Get requested format
        output_format = request.args.get('format', 'json').lower()
        
        if output_format == 'json':
            # Return as JSON
            return {
                "status": "success",
                "data": {
                    "count": len(results),
                    "test_cases": results
                }
            }, 200
            
        elif output_format == 'excel':
            # Create Excel file in memory
            output = BytesIO()
            
            # Convert results to DataFrame
            df = pd.DataFrame(results)
            
            # Write to Excel
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, index=False)
            
            # Prepare for download
            output.seek(0)
            
            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f"test_case_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            )
            
        elif output_format == 'pdf':
            # Create a PDF report
            # In a real implementation, this would generate a more comprehensive PDF
            # For now, just create a simple one with reportlab
            
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            
            # Add title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, height - 50, "Test Case Report")
            
            # Add date
            c.setFont("Helvetica", 10)
            c.drawString(50, height - 70, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Add filter criteria
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, height - 100, "Filter Criteria:")
            
            c.setFont("Helvetica", 10)
            y_pos = height - 120
            for key, value in criteria.items():
                if isinstance(value, dict):
                    c.drawString(50, y_pos, f"{key}: {value['op']} {value['value']}")
                else:
                    c.drawString(50, y_pos, f"{key}: {value}")
                y_pos -= 15
            
            # Add results count
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y_pos - 20, f"Results: {len(results)} test cases")
            
            # Add results table header
            c.setFont("Helvetica-Bold", 10)
            y_pos -= 40
            c.drawString(50, y_pos, "ID")
            c.drawString(150, y_pos, "Test Case")
            c.drawString(300, y_pos, "Owner")
            c.drawString(400, y_pos, "Status")
            c.drawString(500, y_pos, "Module")
            
            # Add line below header
            y_pos -= 10
            c.line(50, y_pos, 550, y_pos)
            y_pos -= 15
            
            # Add results
            c.setFont("Helvetica", 8)
            for result in results[:30]:  # Limit to 30 rows for simplicity
                # Check if we need a new page
                if y_pos < 50:
                    c.showPage()
                    
                    # Reset position and add header
                    y_pos = height - 50
                    c.setFont("Helvetica-Bold", 10)
                    c.drawString(50, y_pos, "ID")
                    c.drawString(150, y_pos, "Test Case")
                    c.drawString(300, y_pos, "Owner")
                    c.drawString(400, y_pos, "Status")
                    c.drawString(500, y_pos, "Module")
                    
                    # Add line below header
                    y_pos -= 10
                    c.line(50, y_pos, 550, y_pos)
                    y_pos -= 15
                    c.setFont("Helvetica", 8)
                
                # Add row
                c.drawString(50, y_pos, str(result.get("TEST_CASE_ID", "")))
                c.drawString(150, y_pos, str(result.get("TEST_CASE", ""))[:20])
                c.drawString(300, y_pos, str(result.get("OWNER", ""))[:15])
                c.drawString(400, y_pos, str(result.get("STATUS", "")))
                c.drawString(500, y_pos, str(result.get("MODULE", ""))[:15])
                
                y_pos -= 12
            
            # If more results, indicate truncation
            if len(results) > 30:
                c.setFont("Helvetica-Oblique", 8)
                c.drawString(50, y_pos - 12, f"... and {len(results) - 30} more test cases (truncated for readability)")
            
            # Save PDF
            c.save()
            buffer.seek(0)
            
            return send_file(
                buffer,
                mimetype='application/pdf',
                as_attachment=True,
                download_name=f"test_case_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            )
            
        else:
            return {"status": "error", "message": f"Unsupported format: {output_format}"}, 400
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/reports/execution-status', methods=['GET'])
def get_execution_status_report() -> Tuple[Dict[str, Any], int]:
    """
    Get a report of test execution status across test cases.
    
    Query parameters:
        - status (str, optional): Filter by status.
        - owner (str, optional): Filter by owner.
        - module (str, optional): Filter by module.
        - executed_after (str, optional): Filter by execution date.
        - executed_before (str, optional): Filter by execution date.
        - result (str, optional): Filter by execution result.
        - format (str, optional): Output format (json, excel, pdf).
        
    Returns:
        Tuple[Dict[str, Any], int]: Report data and HTTP status code.
    """
    try:
        # Collect filter criteria
        criteria = {}
        
        # Add filters from query parameters
        if 'status' in request.args:
            criteria["STATUS"] = request.args.get('status')
        
        if 'owner' in request.args:
            criteria["OWNER"] = request.args.get('owner')
        
        if 'module' in request.args:
            criteria["MODULE"] = request.args.get('module')
        
        if 'result' in request.args:
            criteria["LAST_EXECUTION_RESULT"] = request.args.get('result')
        
        # Date-based filters for execution
        if 'executed_after' in request.args:
            criteria["LAST_EXECUTION_DATE"] = {"op": ">", "value": request.args.get('executed_after')}
        
        if 'executed_before' in request.args:
            if "LAST_EXECUTION_DATE" in criteria:
                # Update existing criteria to a between operation
                executed_after = criteria["LAST_EXECUTION_DATE"]["value"]
                criteria["LAST_EXECUTION_DATE"] = {"op": "between", "value": [executed_after, request.args.get('executed_before')]}
            else:
                criteria["LAST_EXECUTION_DATE"] = {"op": "<", "value": request.args.get('executed_before')}
        
        # Search test cases with criteria
        results = metadata_manager.search_test_cases(criteria)
        
        # Process results for execution status report
        results_by_status = {}
        for result in results:
            status = result.get("LAST_EXECUTION_RESULT", "Not Executed")
            if status not in results_by_status:
                results_by_status[status] = []
            results_by_status[status].append(result)
        
        # Calculate summary
        summary = {
            "total": len(results),
            "by_result": {status: len(items) for status, items in results_by_status.items()},
            "execution_rate": 0
        }
        
        # Calculate execution rate (percent of test cases executed)
        executed_count = sum(len(items) for status, items in results_by_status.items() if status != "Not Executed")
        summary["execution_rate"] = (executed_count / len(results)) * 100 if len(results) > 0 else 0
        
        # Get requested format
        output_format = request.args.get('format', 'json').lower()
        
        if output_format == 'json':
            # Return as JSON
            return {
                "status": "success",
                "data": {
                    "summary": summary,
                    "results_by_status": results_by_status
                }
            }, 200
            
        elif output_format == 'excel':
            # Create Excel file in memory
            output = BytesIO()
            
            # Create a workbook with multiple sheets
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                # Summary sheet
                summary_df = pd.DataFrame({
                    "Metric": ["Total Test Cases", "Execution Rate"] + [f"{status} Count" for status in results_by_status.keys()],
                    "Value": [summary["total"], f"{summary['execution_rate']:.1f}%"] + [len(items) for items in results_by_status.values()]
                })
                summary_df.to_excel(writer, sheet_name="Summary", index=False)
                
                # Sheet for each status
                for status, items in results_by_status.items():
                    if items:
                        status_df = pd.DataFrame(items)
                        status_sheet_name = status.replace(" ", "_")[:31]  # Excel sheet name limit is 31 chars
                        status_df.to_excel(writer, sheet_name=status_sheet_name, index=False)
            
            # Prepare for download
            output.seek(0)
            
            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f"execution_status_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            )
            
        else:
            return {"status": "error", "message": f"Unsupported format: {output_format}"}, 400
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/database/maintenance', methods=['POST'])
def perform_database_maintenance() -> Tuple[Dict[str, Any], int]:
    """
    Perform database maintenance operations.
    
    Request body:
        - operation (str): The maintenance operation to perform.
        - confirmation (str): Confirmation code (required for some operations).
        
    Returns:
        Tuple[Dict[str, Any], int]: Result and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'operation' not in request_data:
            return {"status": "error", "message": "Missing operation"}, 400
        
        operation = request_data['operation']
        confirmation = request_data.get('confirmation')
        
        # Perform the requested operation
        if operation == 'vacuum':
            # Vacuum database to reclaim storage and update statistics
            result = metadata_manager.vacuum_database()
            
            return {
                "status": "success" if result else "error",
                "message": "Database vacuum completed successfully" if result else "Database vacuum failed",
                "data": {
                    "operation": operation,
                    "timestamp": datetime.now().isoformat()
                }
            }, 200 if result else 500
            
        elif operation == 'cleanup':
            # Clean up orphaned data
            result = metadata_manager.cleanup_orphaned_data()
            
            return {
                "status": "success",
                "message": "Database cleanup completed",
                "data": {
                    "operation": operation,
                    "cleaned_items": result,
                    "timestamp": datetime.now().isoformat()
                }
            }, 200
            
        elif operation == 'backup':
            # Create a database backup
            
            # Check for output path
            if 'output_path' not in request_data:
                return {"status": "error", "message": "Missing output_path for backup operation"}, 400
            
            output_path = request_data['output_path']
            include_files = request_data.get('include_files', True)
            
            # Create backup
            result = metadata_manager.create_backup(output_path, include_files)
            
            return {
                "status": "success",
                "message": "Database backup completed",
                "data": {
                    "operation": operation,
                    "backup_stats": result,
                    "output_path": output_path,
                    "timestamp": datetime.now().isoformat()
                }
            }, 200
            
        elif operation == 'restore':
            # Restore from a database backup
            
            # Check for backup path
            if 'backup_path' not in request_data:
                return {"status": "error", "message": "Missing backup_path for restore operation"}, 400
            
            backup_path = request_data['backup_path']
            
            # Check for confirmation
            if not confirmation or confirmation != 'CONFIRM_RESTORE':
                return {"status": "error", "message": "Restore operation requires confirmation code 'CONFIRM_RESTORE'"}, 400
            
            # Restore from backup
            result = metadata_manager.restore_from_backup(backup_path)
            
            return {
                "status": "success",
                "message": "Database restore completed",
                "data": {
                    "operation": operation,
                    "restore_stats": result,
                    "backup_path": backup_path,
                    "timestamp": datetime.now().isoformat()
                }
            }, 200
            
        else:
            return {"status": "error", "message": f"Unsupported operation: {operation}"}, 400
        
    except Exception as e:
        return handle_error(e)
    

@app.route('/api/requirements/processed', methods=['GET'])
def get_processed_requirements() -> Tuple[Dict[str, Any], int]:
    """
    Get processed requirements from the system.
    
    Returns:
        Tuple[Dict[str, Any], int]: Requirements data and HTTP status code.
    """
    try:
        # In a real implementation, this would fetch from a database or the LLM test scenario generator
        # For now, we'll return mock data for demonstration
        
        # Fetch from llm_test_scenario_generator module if available
        try:
            from src.phase1.llm_test_scenario_generator.scenario_generator import ScenarioGenerator
            scenario_generator = ScenarioGenerator()
            requirements = []
            # In the future, this would call a method like scenario_generator.get_processed_requirements()
        except (ImportError, AttributeError):
            # Fallback to mock data if the module is not available
            requirements = [
                {
                    "id": "REQ-001",
                    "title": "User Authentication",
                    "description": "The system shall provide user authentication functionality",
                    "source": "JIRA",
                    "status": "Processed",
                    "priority": "High"
                },
                {
                    "id": "REQ-002",
                    "title": "Password Reset",
                    "description": "Users shall be able to reset their passwords",
                    "source": "JIRA",
                    "status": "Processed",
                    "priority": "Medium"
                }
            ]
        
        return {
            "status": "success",
            "data": {
                "requirements": requirements
            }
        }, 200
    except Exception as e:
        return handle_error(e)

@app.route('/api/requirements/upload', methods=['POST'])
def upload_requirements() -> Tuple[Dict[str, Any], int]:
    """
    Upload requirements from various sources (file, text, JIRA).
    
    Request body:
        - source_type (str): Type of source (file, text, jira).
        - content (Any): The content to process.
        
    Returns:
        Tuple[Dict[str, Any], int]: Upload result and HTTP status code.
    """
    try:
        # Get the source type
        source_type = request.form.get('source_type') if request.form else request.json.get('source_type') if request.json else None
        
        if not source_type:
            return {"status": "error", "message": "Missing source_type"}, 400
        
        # Process based on source type
        if source_type == 'file':
            # Process file upload
            if 'file' not in request.files:
                return {"status": "error", "message": "No file provided"}, 400
                
            file = request.files['file']
            if not file.filename:
                return {"status": "error", "message": "Empty file"}, 400
                
            # Get file extension
            file_ext = os.path.splitext(file.filename)[1].lower()
            
            # Create a temporary file
            with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as temp_file:
                file.save(temp_file.name)
                temp_path = temp_file.name
            
            try:
                # Process requirements from file
                # In a real implementation, this would call a method to process the file
                # For now, just simulate with a response
                
                # Read file content (just for demonstration)
                if file_ext in ['.docx', '.doc']:
                    try:
                        doc = Document(temp_path)
                        num_paragraphs = len(doc.paragraphs)
                    except Exception:
                        num_paragraphs = 0
                elif file_ext in ['.xlsx', '.xls']:
                    try:
                        df = pd.read_excel(temp_path)
                        num_rows = len(df)
                    except Exception:
                        num_rows = 0
                else:
                    # For other file types, just get size
                    file_size = os.path.getsize(temp_path)
                
                # Generate mock requirements
                requirements = []
                if file_ext in ['.docx', '.doc']:
                    for i in range(min(5, num_paragraphs)):
                        requirements.append({
                            "id": f"REQ-{1000+i}",
                            "title": f"Requirement from {file.filename} - {i+1}",
                            "description": f"Requirement extracted from document paragraph {i+1}",
                            "source": "Document Upload",
                            "status": "Processed",
                            "priority": "Medium"
                        })
                elif file_ext in ['.xlsx', '.xls']:
                    for i in range(min(5, num_rows)):
                        requirements.append({
                            "id": f"REQ-{1000+i}",
                            "title": f"Requirement from {file.filename} - {i+1}",
                            "description": f"Requirement extracted from Excel row {i+1}",
                            "source": "Excel Upload",
                            "status": "Processed",
                            "priority": "Medium"
                        })
                else:
                    # Generic file
                    for i in range(3):
                        requirements.append({
                            "id": f"REQ-{1000+i}",
                            "title": f"Requirement from {file.filename} - {i+1}",
                            "description": f"Requirement extracted from file with size {file_size} bytes",
                            "source": "File Upload",
                            "status": "Processed",
                            "priority": "Medium"
                        })
                
                return {
                    "status": "success",
                    "message": f"Processed requirements from {file.filename}",
                    "data": {
                        "source_type": source_type,
                        "source": file.filename,
                        "requirements": requirements
                    }
                }, 200
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
        
        elif source_type == 'text':
            # Process text input
            text_content = request.form.get('content') if request.form else request.json.get('content')
            
            if not text_content:
                return {"status": "error", "message": "Missing content"}, 400
            
            # In a real implementation, this would parse the text and extract requirements
            # For now, just simulate with a response
            
            # Split by lines as a simple extraction method
            lines = text_content.strip().split('\n')
            requirements = []
            
            for i, line in enumerate(lines[:5]):  # Process up to 5 lines
                if line.strip():  # Skip empty lines
                    requirements.append({
                        "id": f"REQ-{2000+i}",
                        "title": f"Requirement from text - {i+1}",
                        "description": line.strip(),
                        "source": "Manual Input",
                        "status": "Processed",
                        "priority": "Medium"
                    })
            
            return {
                "status": "success",
                "message": "Processed requirements from text input",
                "data": {
                    "source_type": source_type,
                    "source": "Text Input",
                    "requirements": requirements
                }
            }, 200
        
        elif source_type == 'jira':
            # Process JIRA input
            jira_project = request.form.get('jira_project') if request.form else request.json.get('jira_project')
            jira_query = request.form.get('jira_query') if request.form else request.json.get('jira_query')
            
            if not jira_project and not jira_query:
                return {"status": "error", "message": "Missing JIRA project or query"}, 400
            
            # In a real implementation, this would connect to JIRA and fetch issues
            # For now, just simulate with a response
            requirements = []
            
            # Generate mock requirements from JIRA
            for i in range(5):
                requirements.append({
                    "id": f"JIRA-{3000+i}",
                    "title": f"Requirement from JIRA - {i+1}",
                    "description": f"Requirement fetched from JIRA project {jira_project or 'N/A'} with query {jira_query or 'N/A'}",
                    "source": "JIRA",
                    "status": "Processed",
                    "priority": "High" if i < 2 else "Medium"
                })
            
            return {
                "status": "success",
                "message": f"Processed requirements from JIRA {jira_project or jira_query}",
                "data": {
                    "source_type": source_type,
                    "source": f"JIRA - {jira_project or jira_query}",
                    "requirements": requirements
                }
            }, 200
        
        else:
            return {"status": "error", "message": f"Unsupported source type: {source_type}"}, 400
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/requirements/process', methods=['POST'])
def process_requirements() -> Tuple[Dict[str, Any], int]:
    """
    Process requirements to generate test scenarios.
    
    Request body:
        - requirements (List[Dict]): List of requirements to process.
        - processing_options (Dict, optional): Options for processing.
        
    Returns:
        Tuple[Dict[str, Any], int]: Processing result and HTTP status code.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'requirements' not in request_data:
            return {"status": "error", "message": "Missing requirements"}, 400
        
        requirements = request_data['requirements']
        processing_options = request_data.get('processing_options', {})
        
        # In a real implementation, this would call LLM to process requirements into test scenarios
        # For now, simulate with a basic transformation
        
        # Generate mock scenarios
        scenarios = []
        for i, req in enumerate(requirements):
            req_id = req.get('id', f"REQ-{i}")
            req_title = req.get('title', f"Requirement {i+1}")
            req_description = req.get('description', "")
            
            # Create one or more scenarios for each requirement
            num_scenarios = 1 + (i % 3)  # 1-3 scenarios per requirement
            
            for j in range(num_scenarios):
                scenario = {
                    "id": f"{req_id}-SC-{j+1}",
                    "name": f"Test scenario for {req_title} - {j+1}",
                    "description": f"Verify {req_description}",
                    "requirement_id": req_id,
                    "priority": req.get('priority', "Medium"),
                    "steps": []
                }
                
                # Add basic steps to the scenario
                scenario["steps"].append({
                    "step_no": 1,
                    "description": "Setup test environment",
                    "expected_result": "Environment is ready for testing"
                })
                
                scenario["steps"].append({
                    "step_no": 2,
                    "description": f"Perform action for {req_title}",
                    "expected_result": "Action is performed successfully"
                })
                
                scenario["steps"].append({
                    "step_no": 3,
                    "description": f"Verify result for {req_title}",
                    "expected_result": "Result is as expected"
                })
                
                scenarios.append(scenario)
        
        # Return the scenarios
        return {
            "status": "success",
            "message": f"Generated {len(scenarios)} test scenarios from {len(requirements)} requirements",
            "data": {
                "scenarios": scenarios,
                "options_applied": processing_options
            }
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/jira/integration/status', methods=['GET'])
def get_jira_integration_status() -> Tuple[Dict[str, Any], int]:
    """
    Get the status of JIRA integration.
    
    Returns:
        Tuple[Dict[str, Any], int]: Integration status and HTTP status code.
    """
    try:
        # In a real implementation, this would check the JIRA connection
        # For now, just return a mock status
        
        jira_status = {
            "connected": True,
            "server_url": os.environ.get("JIRA_URL", "https://example.atlassian.net"),
            "projects_available": ["PROJECT1", "PROJECT2", "PROJECT3"],
            "last_sync": datetime.now().isoformat(),
            "issues_synced": 250,
            "authentication_method": "API Token",
            "user": os.environ.get("JIRA_USER", "jira.user@example.com")
        }
        
        return {
            "status": "success",
            "data": jira_status
        }, 200
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/sharepoint/integration/status', methods=['GET'])
def get_sharepoint_integration_status() -> Tuple[Dict[str, Any], int]:
    """
    Get the status of SharePoint integration.
    
    Returns:
        Tuple[Dict[str, Any], int]: Integration status and HTTP status code.
    """
    try:
        # In a real implementation, this would check the SharePoint connection
        # For now, just return a mock status
        
        sharepoint_status = {
            "connected": True,
            "site_url": os.environ.get("SHAREPOINT_URL", "https://example.sharepoint.com/sites/testcases"),
            "libraries_available": ["Test Cases", "Requirements", "Test Results"],
            "last_sync": datetime.now().isoformat(),
            "files_synced": 125,
            "authentication_method": "Microsoft Graph API",
            "user": os.environ.get("SHAREPOINT_USER", "sharepoint.user@example.com")
        }
        
        return {
            "status": "success",
            "data": sharepoint_status
        }, 200
        
    except Exception as e:
        return handle_error(e)
    

@app.route('/api/test-cases/generate-download', methods=['POST'])
def generate_test_cases_download() -> Response:
    """
    Generate test cases for download in specified format.
    
    Request body:
        - scenarios (List[Dict]): Test scenarios
        - format (str): Output format (excel, word, pdf)
        
    Returns:
        Response: File download response.
    """
    try:
        # Get request data
        request_data = request.json
        
        if not request_data or 'scenarios' not in request_data:
            return jsonify({"status": "error", "message": "Missing scenarios data"}), 400
        
        scenarios = request_data['scenarios']
        output_format = request_data.get('format', 'excel').lower()
        
        # Generate test cases
        test_cases = {}
        for scenario in scenarios:
            try:
                test_case_df = test_case_generator.generate_test_case_from_scenario(scenario)
                test_cases[scenario.get('id', str(uuid.uuid4()))] = test_case_df
            except Exception as scenario_error:
                logger.error(f"Error generating test case for scenario {scenario.get('id', 'unknown')}: {str(scenario_error)}")
                # Continue with other scenarios
        
        if not test_cases:
            return jsonify({"status": "error", "message": "Failed to generate any test cases"}), 500
        
        # Create appropriate format for download
        if output_format == 'excel' or output_format == 'xlsx':
            # Create Excel file in memory
            output = BytesIO()
            
            # Save to Excel with multiple sheets
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                for scenario_id, test_case_df in test_cases.items():
                    # Truncate scenario_id to avoid Excel sheet name limit
                    sheet_name = f"TC_{scenario_id}"
                    if len(sheet_name) > 31:  # Excel sheet name limit is 31 chars
                        sheet_name = sheet_name[:31]
                    test_case_df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            # Prepare for download
            output.seek(0)
            
            # Return the file for download
            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f"test_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            )
            
        elif output_format == 'word' or output_format == 'docx':
            # Create Word document using python-docx
            document = Document()
            
            for scenario_id, test_case_df in test_cases.items():
                # Get basic test case info from first row
                if len(test_case_df) > 0:
                    first_row = test_case_df.iloc[0]
                    title = first_row.get('TEST CASE', 'Test Case')
                    test_case_id = first_row.get('TEST CASE NUMBER', scenario_id)
                    subject = first_row.get('SUBJECT', 'Unknown')
                else:
                    title = 'Test Case'
                    test_case_id = scenario_id
                    subject = 'Unknown'
                
                # Add test case header
                document.add_heading(f"{title} ({test_case_id})", level=1)
                document.add_paragraph(f"Subject: {subject}")
                
                # Add test steps table
                table = document.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Step'
                header_cells[1].text = 'Description'
                header_cells[2].text = 'Test Data'
                header_cells[3].text = 'Expected Result'
                
                # Add steps
                for _, row in test_case_df.iterrows():
                    new_row = table.add_row().cells
                    new_row[0].text = str(row.get('STEP NO', ''))
                    new_row[1].text = str(row.get('TEST STEP DESCRIPTION', ''))
                    new_row[2].text = str(row.get('DATA', ''))
                    new_row[3].text = str(row.get('EXPECTED RESULT', ''))
                
                # Add page break between test cases
                document.add_page_break()
            
            # Save document to memory
            output = BytesIO()
            document.save(output)
            output.seek(0)
            
            # Return the file for download
            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f"test_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            
        elif output_format == 'pdf':
            # Create a PDF using reportlab
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            margin = 50
            
            for scenario_id, test_case_df in test_cases.items():
                # Get basic test case info from first row
                if len(test_case_df) > 0:
                    first_row = test_case_df.iloc[0]
                    title = first_row.get('TEST CASE', 'Test Case')
                    test_case_id = first_row.get('TEST CASE NUMBER', scenario_id)
                    subject = first_row.get('SUBJECT', 'Unknown')
                else:
                    title = 'Test Case'
                    test_case_id = scenario_id
                    subject = 'Unknown'
                
                # Reset position for new page
                y_position = height - margin
                
                # Add test case header
                c.setFont("Helvetica-Bold", 14)
                c.drawString(margin, y_position, f"{title} ({test_case_id})")
                y_position -= 20
                
                c.setFont("Helvetica", 10)
                c.drawString(margin, y_position, f"Subject: {subject}")
                y_position -= 30
                
                # Add table header
                c.setFont("Helvetica-Bold", 10)
                c.drawString(margin, y_position, "Step")
                c.drawString(margin + 40, y_position, "Description")
                c.drawString(margin + 250, y_position, "Test Data")
                c.drawString(margin + 350, y_position, "Expected Result")
                y_position -= 15
                
                c.line(margin, y_position, width - margin, y_position)
                y_position -= 15
                
                # Add steps
                c.setFont("Helvetica", 8)
                for _, row in test_case_df.iterrows():
                    # Check if we need a new page
                    if y_position < margin + 20:
                        c.showPage()
                        
                        # Reset position and add header
                        y_position = height - margin
                        c.setFont("Helvetica-Bold", 10)
                        c.drawString(margin, y_position, "Step")
                        c.drawString(margin + 40, y_position, "Description")
                        c.drawString(margin + 250, y_position, "Test Data")
                        c.drawString(margin + 350, y_position, "Expected Result")
                        y_position -= 15
                        
                        c.line(margin, y_position, width - margin, y_position)
                        y_position -= 15
                        c.setFont("Helvetica", 8)
                    
                    # Add row
                    step_no = str(row.get('STEP NO', ''))
                    description = str(row.get('TEST STEP DESCRIPTION', ''))
                    data = str(row.get('DATA', ''))
                    expected = str(row.get('EXPECTED RESULT', ''))
                    
                    c.drawString(margin, y_position, step_no)
                    
                    # Handle multiline text (simplified approach)
                    wrapped_description = []
                    for i in range(0, len(description), 35):
                        wrapped_description.append(description[i:i+35])
                    
                    for i, line in enumerate(wrapped_description):
                        if i == 0:
                            c.drawString(margin + 40, y_position, line)
                        else:
                            y_position -= 12
                            c.drawString(margin + 40, y_position, line)
                    
                    c.drawString(margin + 250, y_position, data[:20])
                    c.drawString(margin + 350, y_position, expected[:30])
                    
                    y_position -= 20
                
                # Add page break between test cases
                c.showPage()
            
            # Save PDF
            c.save()
            buffer.seek(0)
            
            # Return the file for download
            return send_file(
                buffer,
                mimetype='application/pdf',
                as_attachment=True,
                download_name=f"test_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            )
            
        else:
            return jsonify({"status": "error", "message": f"Unsupported format: {output_format}"}), 400
            
    except Exception as e:
        error_response, status_code = handle_error(e)
        return jsonify(error_response), status_code

@app.route('/api/utility/convert-excel-to-test-case', methods=['POST'])
def convert_excel_to_test_case() -> Tuple[Dict[str, Any], int]:
    """
    Convert a custom Excel file to a test case format.
    
    Request body for multipart/form-data:
        - file: Excel file to convert.
        
    Returns:
        Tuple[Dict[str, Any], int]: Conversion result and HTTP status code.
    """
    try:
        # Check if we have a file
        if 'file' not in request.files:
            return {"status": "error", "message": "No file provided"}, 400
            
        file = request.files['file']
        if not file.filename:
            return {"status": "error", "message": "Empty file"}, 400
            
        # Check file extension
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in ['.xlsx', '.xls']:
            return {"status": "error", "message": "Only Excel files (.xlsx, .xls) are supported"}, 400
            
        # Create a temporary file
        with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as temp_file:
            file.save(temp_file.name)
            temp_path = temp_file.name
            
        try:
            # Read Excel file
            df = pd.read_excel(temp_path)
            
            # Check if it's already in test case format
            required_columns = ["TEST CASE NUMBER", "TEST STEP DESCRIPTION", "EXPECTED RESULT"]
            is_test_case_format = all(col in df.columns for col in required_columns)
            
            if is_test_case_format:
                # Already in test case format, just return it
                result = {
                    "status": "success",
                    "message": "File is already in test case format",
                    "data": {
                        "test_case": df.to_dict('records'),
                        "already_formatted": True
                    }
                }
            else:
                # Try to convert to test case format
                # This is a simple example - in a real implementation, this would be more sophisticated
                
                # Create new DataFrame in test case format
                test_case_df = pd.DataFrame(columns=[
                    "SUBJECT", "TEST CASE", "TEST CASE NUMBER", "STEP NO", 
                    "TEST STEP DESCRIPTION", "DATA", "EXPECTED RESULT", 
                    "TEST USER ID/ROLE", "STATUS"
                ])
                
                # Fill in with data from source
                test_case_number = f"TC-{uuid.uuid4().hex[:8].upper()}"
                test_case_name = os.path.splitext(file.filename)[0]
                
                # Try to find columns that might contain steps and expected results
                step_col = None
                result_col = None
                
                # Look for columns with names that might contain steps
                for col in df.columns:
                    col_lower = str(col).lower()
                    if any(term in col_lower for term in ['step', 'action', 'procedure', 'instruction']):
                        step_col = col
                    elif any(term in col_lower for term in ['result', 'expected', 'outcome', 'verification']):
                        result_col = col
                
                if step_col and result_col:
                    # We found columns that might contain steps and results
                    for i, row in df.iterrows():
                        test_case_df.loc[i] = [
                            "Converted", # SUBJECT
                            test_case_name, # TEST CASE
                            test_case_number, # TEST CASE NUMBER
                            i + 1, # STEP NO
                            str(row[step_col]), # TEST STEP DESCRIPTION
                            "", # DATA
                            str(row[result_col]), # EXPECTED RESULT
                            "Tester", # TEST USER ID/ROLE
                            "Draft" # STATUS
                        ]
                else:
                    # No matching columns found, make a best guess
                    # Use the first column as steps and second as results if available
                    for i, row in df.iterrows():
                        test_case_df.loc[i] = [
                            "Converted", # SUBJECT
                            test_case_name, # TEST CASE
                            test_case_number, # TEST CASE NUMBER
                            i + 1, # STEP NO
                            str(row.iloc[0]) if df.shape[1] > 0 else f"Step {i+1}", # TEST STEP DESCRIPTION
                            "", # DATA
                            str(row.iloc[1]) if df.shape[1] > 1 else "", # EXPECTED RESULT
                            "Tester", # TEST USER ID/ROLE
                            "Draft" # STATUS
                        ]
                
                result = {
                    "status": "success",
                    "message": "File converted to test case format",
                    "data": {
                        "test_case": test_case_df.to_dict('records'),
                        "already_formatted": False
                    }
                }
            
            return result, 200
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
    except Exception as e:
        return handle_error(e)

@app.route('/api/utility/health-check', methods=['GET'])
def utility_health_check() -> Tuple[Dict[str, Any], int]:
    """
    Perform a comprehensive health check of the system.
    
    Returns:
        Tuple[Dict[str, Any], int]: Health check results and HTTP status code.
    """
    try:
        # Check API health
        api_status = {
            "status": "healthy",
            "uptime": "Unknown"  # In a real implementation, this would track server uptime
        }
        
        # Check database health
        db_status = "unknown"
        try:
            db_status = "connected" if metadata_manager.check_database_connection() else "disconnected"
        except Exception as e:
            logger.error(f"Database health check failed: {str(e)}")
            db_status = "error"
        
        # Check components
        components_status = {
            "test_case_generator": test_case_generator is not None,
            "test_case_refiner": test_case_refiner is not None,
            "version_controller": version_controller is not None,
            "metadata_manager": metadata_manager is not None
        }
        
        # Check integrations
        integrations = {
            "jira": {
                "status": "unknown",
                "message": "JIRA integration status check not implemented"
            },
            "sharepoint": {
                "status": "unknown",
                "message": "SharePoint integration status check not implemented"
            },
            "alm": {
                "status": "unknown",
                "message": "ALM integration status check not implemented"
            }
        }
        
        # For a real implementation, check each integration
        # For now, we'll use mock data
        integrations["jira"]["status"] = "connected"
        integrations["sharepoint"]["status"] = "connected"
        integrations["alm"]["status"] = "disconnected"
        
        # Check disk space
        disk_space = {
            "available": "Unknown",
            "used": "Unknown",
            "total": "Unknown"
        }
        
        try:
            import shutil
            total, used, free = shutil.disk_usage('/')
            
            # Convert to human-readable format
            def human_readable_size(size_bytes):
                for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
                    if size_bytes < 1024.0 or unit == 'TB':
                        return f"{size_bytes:.2f} {unit}"
                    size_bytes /= 1024.0
            
            disk_space = {
                "available": human_readable_size(free),
                "used": human_readable_size(used),
                "total": human_readable_size(total),
                "percent_used": f"{used/total*100:.1f}%"
            }
        except Exception as disk_error:
            logger.error(f"Disk space check failed: {str(disk_error)}")
            disk_space["error"] = str(disk_error)
        
        # Combine all health data
        health_data = {
            "api": api_status,
            "database": {
                "status": db_status
            },
            "components": components_status,
            "integrations": integrations,
            "disk_space": disk_space,
            "timestamp": datetime.now().isoformat()
        }
        
        # Determine overall status
        critical_components = ["api", "database"]
        overall_status = "healthy"
        
        for component in critical_components:
            if component == "api" and health_data[component]["status"] != "healthy":
                overall_status = "unhealthy"
                break
            if component == "database" and health_data[component]["status"] != "connected":
                overall_status = "unhealthy"
                break
        
        health_data["overall_status"] = overall_status
        
        return {
            "status": "success",
            "data": health_data
        }, 200 if overall_status == "healthy" else 503
        
    except Exception as e:
        return handle_error(e)

@app.route('/api/documentation', methods=['GET'])
def get_api_documentation() -> Tuple[Dict[str, Any], int]:
    """
    Get API documentation for all endpoints.
    
    Returns:
        Tuple[Dict[str, Any], int]: API documentation and HTTP status code.
    """
    try:
        # Create documentation for all endpoints
        endpoints = []
        
        for rule in app.url_map.iter_rules():
            if rule.endpoint != 'static':
                endpoint = {
                    "path": str(rule),
                    "methods": list(rule.methods - {'HEAD', 'OPTIONS'}),
                    "endpoint": rule.endpoint
                }
                
                # Get function for the endpoint
                func = app.view_functions[rule.endpoint]
                
                # Get docstring if available
                if func.__doc__:
                    docstring = func.__doc__.strip()
                    # Parse docstring to extract description and parameters
                    lines = docstring.split('\n')
                    description = lines[0].strip()
                    params = []
                    returns = []
                    
                    current_section = None
                    for line in lines[1:]:
                        line = line.strip()
                        if not line:
                            continue
                            
                        if line.startswith('Request body:'):
                            current_section = 'params'
                        elif line.startswith('Args:'):
                            current_section = 'params'
                        elif line.startswith('Query parameters:'):
                            current_section = 'params'
                        elif line.startswith('Returns:'):
                            current_section = 'returns'
                        elif current_section == 'params' and line.startswith('-'):
                            params.append(line.lstrip('- '))
                        elif current_section == 'returns' and line.startswith('-'):
                            returns.append(line.lstrip('- '))
                    
                    endpoint["description"] = description
                    endpoint["parameters"] = params
                    endpoint["returns"] = returns
                
                endpoints.append(endpoint)
        
        # Group endpoints by category
        categories = {}
        
        for endpoint in endpoints:
            path = endpoint["path"]
            category = path.split('/')[2] if len(path.split('/')) > 2 else 'other'
            
            if category not in categories:
                categories[category] = []
                
            categories[category].append(endpoint)
        
        # Return documentation
        return {
            "status": "success",
            "data": {
                "api_name": "Test Case API Service",
                "version": "1.0.0",
                "base_url": "/api",
                "categories": categories
            }
        }, 200
        
    except Exception as e:
        return handle_error(e)

# Helper function to convert test case DataFrame to scenario object
def convert_test_case_to_scenario(test_case_df):
    """
    Convert a test case DataFrame to a scenario object
    
    Args:
        test_case_df (pd.DataFrame): Test case DataFrame
        
    Returns:
        Dict: Scenario object
    """
    if len(test_case_df) == 0:
        return None
        
    # Extract basic info from first row
    first_row = test_case_df.iloc[0]
    
    scenario = {
        "id": first_row.get("TEST CASE NUMBER", f"TC-{uuid.uuid4().hex[:8].upper()}"),
        "name": first_row.get("TEST CASE", "Test Case"),
        "subject": first_row.get("SUBJECT", "Unknown"),
        "type": first_row.get("TYPE", "Functional"),
        "steps": []
    }
    
    # Extract steps
    for _, row in test_case_df.iterrows():
        step = {
            "step_no": row.get("STEP NO", 1),
            "description": row.get("TEST STEP DESCRIPTION", ""),
            "data": row.get("DATA", ""),
            "reference_values": row.get("REFERENCE VALUES", ""),
            "values": row.get("VALUES", ""),
            "expected_result": row.get("EXPECTED RESULT", ""),
            "trans_code": row.get("TRANS CODE", ""),
            "test_user": row.get("TEST USER ID/ROLE", "")
        }
        scenario["steps"].append(step)
        
    return scenario

# Main entry point
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Run the Test Case API Service")
    parser.add_argument("--host", default="0.0.0.0", help="Host to run the server on")
    parser.add_argument("--port", type=int, default=5000, help="Port to run the server on")
    parser.add_argument("--debug", action="store_true", help="Run in debug mode")
    
    args = parser.parse_args()
    
    # Configure logging
    logging.basicConfig(
        level=logging.DEBUG if args.debug else logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # Initialize components before starting the server
    if not initialize_components():
        logger.error("Failed to initialize components. Exiting.")
        exit(1)
    
    # Run the Flask app
    app.run(host=args.host, port=args.port, debug=args.debug)

    