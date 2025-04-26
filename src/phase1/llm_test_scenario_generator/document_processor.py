"""
Document Processor Module for the LLM Test Scenario Generator.

This module processes input documents (Word, Excel, PDF, Text) to extract requirements.
It supports various file formats and is designed to identify and structure requirements
for further processing by the LLM-based scenario generator.

Usage:
    processor = DocumentProcessor()
    requirements = processor.process_document('path/to/document.docx')
"""

import os
import re
import io
import logging
from typing import Dict, List, Any, Union, Optional, Tuple

# Third-party libraries for document processing
try:
    import docx  # For Word documents
    import pandas as pd  # For Excel files
    import PyPDF2  # For PDF files
    import mammoth  # For Word documents (alternative to python-docx with better formatting support)
except ImportError:
    logging.warning("Some document processing libraries not found. Please install required dependencies.")

class DocumentProcessor:
    """Class for processing various document types to extract requirements."""

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize the DocumentProcessor.

        Args:
            config: Configuration dictionary with processing settings.
        """
        self.config = config or {}
        self.logger = logging.getLogger(__name__)
        self.setup_logging()

    def setup_logging(self) -> None:
        """Configure logging for the document processor."""
        log_level = self.config.get('log_level', logging.INFO)
        self.logger.setLevel(log_level)
        
        # Create console handler if none exists
        if not self.logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            self.logger.addHandler(handler)

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract requirements.

        Args:
            file_path: Path to the document file.

        Returns:
            Dictionary containing extracted requirements data.
        """
        self.logger.info(f"Processing document: {file_path}")

        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension in ['.docx', '.doc']:
                return self.process_word_document(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                return self.process_excel_document(file_path)
            elif file_extension == '.pdf':
                return self.process_pdf_document(file_path)
            elif file_extension in ['.txt', '.md', '.json']:
                return self.process_text_document(file_path)
            else:
                self.logger.error(f"Unsupported file format: {file_extension}")
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            self.logger.error(f"Error processing document: {str(e)}")
            raise

    def process_word_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a Word document (.docx, .doc) to extract requirements.

        Args:
            file_path: Path to the Word document.

        Returns:
            Dictionary containing extracted requirements data.
        """
        self.logger.info("Processing Word document")
        try:
            # Try using mammoth first (better with formatting)
            with open(file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                text = result.value
            
            # If mammoth extraction is empty or too limited, fallback to python-docx
            if not text or len(text) < 100:
                doc = docx.Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                text = "\n".join(paragraphs)
                
                # Get text from tables if any
                tables_data = []
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    tables_data.append(table_data)
                
                # Extract structured data if available
                return self._extract_requirements_from_text(text, tables_data)
            
            return self._extract_requirements_from_text(text)
            
        except Exception as e:
            self.logger.error(f"Error processing Word document: {str(e)}")
            raise

    def process_excel_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process an Excel document (.xlsx, .xls) to extract requirements.

        Args:
            file_path: Path to the Excel document.

        Returns:
            Dictionary containing extracted requirements data.
        """
        self.logger.info("Processing Excel document")
        try:
            # Read all sheets from the Excel file
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            requirements_data = {
                "document_type": "excel",
                "file_path": file_path,
                "sheets": {}
            }
            
            # Process each sheet
            for sheet_name, df in excel_data.items():
                self.logger.info(f"Processing sheet: {sheet_name}")
                
                # Try to identify if this is a requirements or user stories sheet
                if self._is_requirements_sheet(df):
                    requirements_data["sheets"][sheet_name] = self._process_requirements_sheet(df)
                else:
                    # Just store the dataframe as dictionary for general use
                    requirements_data["sheets"][sheet_name] = df.to_dict(orient='records')
            
            return requirements_data
            
        except Exception as e:
            self.logger.error(f"Error processing Excel document: {str(e)}")
            raise

    def process_pdf_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a PDF document to extract requirements.

        Args:
            file_path: Path to the PDF document.

        Returns:
            Dictionary containing extracted requirements data.
        """
        self.logger.info("Processing PDF document")
        try:
            text = ""
            with open(file_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            
            return self._extract_requirements_from_text(text)
            
        except Exception as e:
            self.logger.error(f"Error processing PDF document: {str(e)}")
            raise

    def process_text_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a plain text document (.txt, .md, .json) to extract requirements.

        Args:
            file_path: Path to the text document.

        Returns:
            Dictionary containing extracted requirements data.
        """
        self.logger.info("Processing text document")
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return self._extract_requirements_from_text(text)
            
        except Exception as e:
            self.logger.error(f"Error processing text document: {str(e)}")
            raise

    def process_jira_export(self, jira_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Process exported JIRA data to extract requirements.

        Args:
            jira_data: Dictionary containing JIRA issue data.

        Returns:
            Dictionary containing extracted requirements data.
        """
        self.logger.info("Processing JIRA export data")
        try:
            requirements_data = {
                "document_type": "jira",
                "stories": []
            }
            
            # Check if it's a single issue or multiple issues
            if 'issues' in jira_data:
                issues = jira_data['issues']
            else:
                issues = [jira_data]
            
            for issue in issues:
                story = self._extract_user_story_from_jira(issue)
                if story:
                    requirements_data["stories"].append(story)
            
            return requirements_data
            
        except Exception as e:
            self.logger.error(f"Error processing JIRA export: {str(e)}")
            raise

    def process_raw_input(self, text: str) -> Dict[str, Any]:
        """
        Process raw text input (e.g., from UI) to extract requirements.

        Args:
            text: Raw text input.

        Returns:
            Dictionary containing extracted requirements data.
        """
        self.logger.info("Processing raw text input")
        return self._extract_requirements_from_text(text)

    def _extract_user_story_from_jira(self, issue: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract user story information from a JIRA issue.

        Args:
            issue: Dictionary containing JIRA issue data.

        Returns:
            Dictionary containing user story information.
        """
        try:
            fields = issue.get('fields', {})
            
            story = {
                "id": issue.get('key', ''),
                "title": fields.get('summary', ''),
                "description": fields.get('description', ''),
                "type": fields.get('issuetype', {}).get('name', ''),
                "priority": fields.get('priority', {}).get('name', ''),
                "status": fields.get('status', {}).get('name', ''),
                "acceptance_criteria": [],
                "metadata": {}
            }
            
            # Try to extract acceptance criteria from custom fields or description
            description = fields.get('description', '')
            if description:
                # Common patterns for acceptance criteria in JIRA
                ac_patterns = [
                    r"Acceptance Criteria:(.*?)(?:\n\n|\n\Z|$)",
                    r"AC:(.*?)(?:\n\n|\n\Z|$)",
                    r"Acceptance Criteria\s*\n((?:.*\n)*?)(?:\n\n|\n\Z|$)"
                ]
                
                for pattern in ac_patterns:
                    match = re.search(pattern, description, re.IGNORECASE | re.DOTALL)
                    if match:
                        ac_text = match.group(1).strip()
                        # Try to extract numbered/bulleted criteria
                        criteria = self._extract_list_items(ac_text)
                        if criteria:
                            story["acceptance_criteria"] = criteria
                            break
            
            # Extract additional custom fields if available
            for key, value in fields.items():
                if key not in ['summary', 'description', 'issuetype', 'priority', 'status']:
                    if isinstance(value, (str, int, float, bool)) or value is None:
                        story["metadata"][key] = value
            
            return story
            
        except Exception as e:
            self.logger.warning(f"Error extracting user story from JIRA issue: {str(e)}")
            return {}

    def _extract_requirements_from_text(self, text: str, tables_data: List[List[List[str]]] = None) -> Dict[str, Any]:
        """
        Extract requirements from text content.

        Args:
            text: Raw text content.
            tables_data: Optional list of tables data extracted from the document.

        Returns:
            Dictionary containing structured requirements data.
        """
        requirements_data = {
            "document_type": "text",
            "raw_text": text,
            "requirements": [],
            "user_stories": [],
            "acceptance_criteria": []
        }
        
        # Extract user stories (common format: As a... I want... So that...)
        user_story_pattern = r"As\s+a\s+(.*?)\s+I\s+want\s+(.*?)\s+(?:So\s+that|To)\s+(.*?)(?:\n\n|\n\Z|$)"
        user_stories = re.finditer(user_story_pattern, text, re.IGNORECASE | re.DOTALL)
        
        for match in user_stories:
            user_story = {
                "role": match.group(1).strip(),
                "goal": match.group(2).strip(),
                "benefit": match.group(3).strip(),
                "full_text": match.group(0).strip()
            }
            requirements_data["user_stories"].append(user_story)
        
        # Extract requirements (look for numbered/bulleted requirements)
        requirements = []
        
        # Try to extract requirements in different formats
        # 1. Numbered requirements (e.g., "1. The system shall...")
        numbered_req_pattern = r"(?:^|\n)(\d+\.[\s\t]*)(.*?)(?=(?:\n\d+\.[\s\t]*)|$)"
        matches = re.finditer(numbered_req_pattern, text, re.MULTILINE | re.DOTALL)
        for match in matches:
            req_text = match.group(2).strip()
            if len(req_text) > 5:  # Minimum length to be considered a requirement
                requirements.append({
                    "id": match.group(1).strip(),
                    "text": req_text,
                    "type": "functional" if self._is_functional_requirement(req_text) else "non-functional"
                })
        
        # 2. Bulleted requirements (e.g., "• The system shall...")
        bulleted_req_pattern = r"(?:^|\n)([•\-\*][\s\t]*)(.*?)(?=(?:\n[•\-\*][\s\t]*)|$)"
        matches = re.finditer(bulleted_req_pattern, text, re.MULTILINE | re.DOTALL)
        for match in matches:
            req_text = match.group(2).strip()
            if len(req_text) > 5:
                requirements.append({
                    "id": match.group(1).strip(),
                    "text": req_text,
                    "type": "functional" if self._is_functional_requirement(req_text) else "non-functional"
                })
        
        # 3. "Shall" requirements (e.g., "The system shall...")
        shall_req_pattern = r"(?:^|\n)(.*?\s+shall\s+.*?)(?=\n|$)"
        matches = re.finditer(shall_req_pattern, text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            req_text = match.group(1).strip()
            if len(req_text) > 10 and not any(req["text"] == req_text for req in requirements):
                requirements.append({
                    "id": "REQ-" + str(len(requirements) + 1),
                    "text": req_text,
                    "type": "functional" if self._is_functional_requirement(req_text) else "non-functional"
                })
        
        # Extract acceptance criteria
        ac_pattern = r"(?:Acceptance\sCriteria|AC)[:\s]+(.*?)(?=\n\n|\n\Z|$)"
        matches = re.finditer(ac_pattern, text, re.IGNORECASE | re.DOTALL)
        for match in matches:
            ac_text = match.group(1).strip()
            criteria = self._extract_list_items(ac_text)
            if criteria:
                requirements_data["acceptance_criteria"].extend(criteria)
        
        # Process tables if available
        if tables_data:
            for table in tables_data:
                # Try to determine if this is a requirements table
                if self._is_requirements_table(table):
                    table_reqs = self._extract_requirements_from_table(table)
                    if table_reqs:
                        requirements.extend(table_reqs)
        
        requirements_data["requirements"] = requirements
        return requirements_data

    def _extract_list_items(self, text: str) -> List[str]:
        """
        Extract list items from text (numbered or bulleted).

        Args:
            text: Text containing list items.

        Returns:
            List of extracted items.
        """
        items = []
        
        # Try numbered list first
        numbered_pattern = r"(?:^|\n)(\d+\.[\s\t]*)(.*?)(?=(?:\n\d+\.[\s\t]*)|$)"
        matches = re.finditer(numbered_pattern, text, re.MULTILINE | re.DOTALL)
        for match in matches:
            items.append(match.group(2).strip())
        
        # If no numbered list found, try bulleted list
        if not items:
            bulleted_pattern = r"(?:^|\n)([•\-\*][\s\t]*)(.*?)(?=(?:\n[•\-\*][\s\t]*)|$)"
            matches = re.finditer(bulleted_pattern, text, re.MULTILINE | re.DOTALL)
            for match in matches:
                items.append(match.group(2).strip())
        
        # If still no items found, try to split by newlines
        if not items and "\n" in text:
            lines = [line.strip() for line in text.split("\n") if line.strip()]
            if lines:
                items = lines
        
        return items

    def _is_functional_requirement(self, text: str) -> bool:
        """
        Determine if a requirement is functional or non-functional.

        Args:
            text: Requirement text.

        Returns:
            True if functional, False if non-functional.
        """
        # Keywords that often indicate non-functional requirements
        nfr_keywords = [
            "performance", "security", "availability", "reliability", 
            "maintainability", "scalability", "usability", "portability",
            "response time", "throughput", "secure", "load", "stress",
            "recovery", "audit", "log", "backup", "restore", "compliance"
        ]
        
        # Check for non-functional requirement keywords
        for keyword in nfr_keywords:
            if re.search(r'\b' + keyword + r'\b', text, re.IGNORECASE):
                return False
        
        # If no non-functional keywords found, assume it's functional
        return True

    def _is_requirements_sheet(self, df: 'pd.DataFrame') -> bool:
        """
        Determine if an Excel sheet contains requirements data.

        Args:
            df: Pandas DataFrame representing the Excel sheet.

        Returns:
            True if it's a requirements sheet, False otherwise.
        """
        # Check column names for typical requirements-related terms
        req_keywords = [
            "requirement", "req", "user story", "feature", "id", "description", 
            "priority", "acceptance", "criteria", "test", "case"
        ]
        
        column_string = " ".join(str(col).lower() for col in df.columns)
        for keyword in req_keywords:
            if keyword in column_string:
                return True
        
        # If column check fails, try checking first few rows for requirements patterns
        sample_rows = df.head(5).astype(str).values.flatten()
        sample_text = " ".join(sample_rows)
        
        # Check for common requirement indicators
        if (re.search(r'shall', sample_text, re.IGNORECASE) or 
            re.search(r'must', sample_text, re.IGNORECASE) or
            re.search(r'as a .* I want', sample_text, re.IGNORECASE)):
            return True
            
        return False

    def _process_requirements_sheet(self, df: 'pd.DataFrame') -> List[Dict[str, Any]]:
        """
        Process a requirements Excel sheet.

        Args:
            df: Pandas DataFrame representing the requirements sheet.

        Returns:
            List of requirement dictionaries.
        """
        requirements = []
        
        # Try to identify key columns
        columns = {col.lower(): col for col in df.columns}
        
        # Map expected fields to actual column names
        field_mapping = {}
        for expected_field, keywords in {
            "id": ["id", "key", "req id", "requirement id", "story id"],
            "description": ["description", "requirement", "text", "details"],
            "type": ["type", "category", "requirement type"],
            "priority": ["priority", "importance", "severity"],
            "status": ["status", "state"]
        }.items():
            for keyword in keywords:
                for col in columns:
                    if keyword in col:
                        field_mapping[expected_field] = columns[col]
                        break
                if expected_field in field_mapping:
                    break
        
        # Process each row
        for _, row in df.iterrows():
            requirement = {}
            
            # Map fields using identified columns
            for field, column in field_mapping.items():
                if pd.notna(row[column]):
                    requirement[field] = str(row[column]).strip()
            
            # Add any other non-empty fields
            for col in df.columns:
                field = col.lower().replace(" ", "_")
                if field not in requirement and pd.notna(row[col]):
                    requirement[field] = str(row[col]).strip()
            
            # Skip empty requirements
            if "description" in requirement and requirement["description"]:
                # Add derived fields
                if "type" not in requirement and "description" in requirement:
                    requirement["type"] = "functional" if self._is_functional_requirement(requirement["description"]) else "non-functional"
                
                requirements.append(requirement)
                
        return requirements

    def _is_requirements_table(self, table: List[List[str]]) -> bool:
        """
        Determine if a table contains requirements data.

        Args:
            table: List of lists representing table rows and cells.

        Returns:
            True if it's a requirements table, False otherwise.
        """
        if not table or len(table) < 2:  # Need at least header + one row
            return False
            
        # Check header row for requirements-related terms
        header_row = " ".join(str(cell).lower() for cell in table[0])
        req_keywords = ["requirement", "req", "id", "description", "user story", "feature"]
        
        for keyword in req_keywords:
            if keyword in header_row:
                return True
                
        return False

    def _extract_requirements_from_table(self, table: List[List[str]]) -> List[Dict[str, Any]]:
        """
        Extract requirements from a table.

        Args:
            table: List of lists representing table rows and cells.

        Returns:
            List of requirement dictionaries.
        """
        requirements = []
        
        if not table or len(table) < 2:
            return requirements
            
        # Assume first row is header
        header = [str(cell).strip().lower() for cell in table[0]]
        
        # Try to identify key columns
        id_col = -1
        desc_col = -1
        type_col = -1
        priority_col = -1
        
        for i, col_name in enumerate(header):
            if any(keyword in col_name for keyword in ["id", "key", "#"]):
                id_col = i
            elif any(keyword in col_name for keyword in ["description", "requirement", "text"]):
                desc_col = i
            elif any(keyword in col_name for keyword in ["type", "category"]):
                type_col = i
            elif any(keyword in col_name for keyword in ["priority", "importance"]):
                priority_col = i
        
        # If we couldn't identify a description column, can't extract requirements
        if desc_col == -1:
            return requirements
            
        # Process each row
        for row_idx in range(1, len(table)):
            row = table[row_idx]
            
            # Skip rows that are too short
            if len(row) <= desc_col:
                continue
                
            requirement = {
                "text": row[desc_col].strip()
            }
            
            # Skip empty requirements
            if not requirement["text"]:
                continue
                
            # Add ID if available
            if id_col != -1 and id_col < len(row) and row[id_col].strip():
                requirement["id"] = row[id_col].strip()
            else:
                requirement["id"] = f"REQ-{row_idx}"
                
            # Add type if available
            if type_col != -1 and type_col < len(row) and row[type_col].strip():
                requirement["type"] = row[type_col].strip().lower()
            else:
                requirement["type"] = "functional" if self._is_functional_requirement(requirement["text"]) else "non-functional"
                
            # Add priority if available
            if priority_col != -1 and priority_col < len(row) and row[priority_col].strip():
                requirement["priority"] = row[priority_col].strip().lower()
                
            requirements.append(requirement)
                
        return requirements


# Example usage
if __name__ == "__main__":
    processor = DocumentProcessor()
    
    # Example with a Word document
    try:
        requirements = processor.process_document("sample_requirements.docx")
        print(f"Extracted {len(requirements['requirements'])} requirements")
        print(f"Extracted {len(requirements['user_stories'])} user stories")
    except Exception as e:
        print(f"Error: {str(e)}")